"""
generar_guia_agente.py
Genera Guia_AI_Agent_Egakat.docx — versión 2.0
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

SALIDA = Path(r"C:\ClaudeWork\AI_Agent\Guia_AI_Agent_Egakat.docx")

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_color(run, r, g, b):
    run.font.color.rgb = RGBColor(r, g, b)

def add_heading(doc, text, level=1, color=(31, 73, 125)):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(*color)
    return p

def add_tabla(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        cell.text = h
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F497D")
        tcPr.append(shd)
    for r_idx, row in enumerate(rows):
        trow = table.rows[r_idx + 1]
        fill = "DCE6F1" if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row):
            cell = trow.cells[c_idx]
            cell.text = str(val)
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), fill)
            tcPr.append(shd)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    return table

def add_caja(doc, titulo, contenido, fill="EBF3FB"):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(f"{titulo}  ")
    r1.bold = True
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(contenido)
    r2.font.color.rgb = RGBColor(50, 50, 50)
    return p

def add_codigo(doc, texto):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    r = p.add_run(texto)
    r.font.name = "Courier New"
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0, 100, 0)
    return p

def pie_pagina(doc):
    footer = doc.sections[0].footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"Sócrates Cabral  |  Control de Gestión y Mejora Continua  |  Egakat SPA  |  "
        f"{datetime.date.today().strftime('%d/%m/%Y')}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor(128, 128, 128)

# ══════════════════════════════════════════════════════════════════════════════
doc = Document()
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── PORTADA ───────────────────────────────────────────────────────────────────
doc.add_paragraph()
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run("LABORATORIO DE IA LOGÍSTICO")
r.bold = True; r.font.size = Pt(26); set_color(r, 31, 73, 125)

t2 = doc.add_paragraph()
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run("Guía de Interacción — Claude.ai + Claude Code + Agentes")
r2.bold = True; r2.font.size = Pt(18); set_color(r2, 68, 114, 196)

doc.add_paragraph()
t3 = doc.add_paragraph()
t3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = t3.add_run("Cómo trabajar con Claude.ai, Claude Code, los agentes y ask_ai de forma coordinada")
r3.font.size = Pt(12); r3.italic = True; set_color(r3, 80, 80, 80)

doc.add_paragraph()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run(
    f"Sócrates Cabral  |  Control de Gestión y Mejora Continua  |  Egakat SPA\n"
    f"Versión 3.0  —  {datetime.date.today().strftime('%d/%m/%Y')}"
)
doc.add_page_break()

# ── 1. VISIÓN GENERAL ─────────────────────────────────────────────────────────
add_heading(doc, "1. Visión general del sistema — las 4 capas", 1)
doc.add_paragraph(
    "El laboratorio de IA logístico de Egakat SPA tiene CUATRO capas que trabajan en secuencia. "
    "No son herramientas separadas — son una sola arquitectura coordinada. "
    "Cada capa tiene un rol distinto y un momento específico de uso:"
)
doc.add_paragraph()
add_tabla(doc,
    ["Capa", "Herramienta", "Rol", "Cuándo usarla"],
    [
        ["0 — Diseño",
         "Claude.ai\n(navegador web)",
         "El arquitecto. Piensa, planifica y diseña sin tocar código. "
         "No tiene acceso al proyecto — trabaja solo con lo que tú le describes.",
         "Al INICIAR un proyecto o tarea nueva. Antes de escribir una sola línea de código."],
        ["1 — Orquestador",
         "Claude Code\n(VS Code)",
         "El cerebro de ejecución. Lee el proyecto completo, decide qué agente llamar, "
         "ejecuta código, interpreta resultados y te entrega el output final.",
         "Para implementar, debugear y construir. Trabaja sobre el código real."],
        ["2 — Agentes",
         "extractor / m365 /\nanalista / generador /\npower_bi",
         "Los especialistas. Cada uno ejecuta una tarea concreta: extraer datos, "
         "analizar, generar scripts, interactuar con OneDrive o crear DAX.",
         "Claude Code los llama automáticamente. También puedes llamarlos tú desde terminal."],
        ["3 — CLI rápido",
         "ask_ai",
         "El asistente de terminal. Consultas puntuales sin abrir VS Code ni una sesión completa.",
         "Cuando necesitas una respuesta rápida sobre un archivo o pregunta específica."],
    ],
    col_widths=[3, 3.5, 8, 4.5]
)
doc.add_paragraph()
add_caja(doc, "Regla de oro:",
    "Piensa con Claude.ai  →  Ejecuta con Claude Code  →  Consulta rápida con ask_ai")
doc.add_paragraph()
doc.add_page_break()

# ── 2. CLAUDE.AI — EL ARQUITECTO ─────────────────────────────────────────────
add_heading(doc, "2. Claude.ai — el arquitecto del sistema", 1)
doc.add_paragraph(
    "Claude.ai (claude.ai en el navegador) es la capa más estratégica del laboratorio. "
    "Es donde nace cada proyecto. No ejecuta código ni tiene acceso a tus archivos, "
    "pero tiene algo que las otras herramientas no tienen: conversación profunda y continua "
    "para pensar problemas complejos sin límite de contexto técnico."
)
doc.add_paragraph()

add_heading(doc, "Qué puede hacer Claude.ai que las otras capas no pueden", 2, color=(68,114,196))
add_tabla(doc,
    ["Capacidad", "Claude.ai", "Claude Code", "ask_ai"],
    [
        ["Diseñar arquitectura de un proyecto",            "✅ Ideal",    "⚠️ Puede",   "❌ No"],
        ["Conversar en profundidad sin código",            "✅ Ideal",    "⚠️ Puede",   "❌ No"],
        ["Revisar lógica de negocio en lenguaje natural",  "✅ Ideal",    "⚠️ Puede",   "❌ No"],
        ["Acceder a archivos del proyecto",                "❌ No",       "✅ Sí",       "✅ Sí"],
        ["Llamar a los agentes automáticamente",           "❌ No",       "✅ Sí",       "❌ No"],
        ["Ejecutar código y ver resultados",               "❌ No",       "✅ Sí",       "✅ Parcial"],
        ["Respuesta rápida desde terminal",                "❌ No",       "❌ No",       "✅ Ideal"],
        ["Mantener contexto entre sesiones (memoria)",     "⚠️ Limitado", "✅ MEMORY.md","❌ No"],
    ],
    col_widths=[7.5, 3, 3.5, 3]
)
doc.add_paragraph()

add_heading(doc, "Cómo usar Claude.ai en tu flujo diario", 2, color=(68,114,196))
doc.add_paragraph(
    "Claude.ai es tu interlocutor para todo lo que requiere pensar antes de hacer. "
    "Úsalo para estas situaciones:"
)
add_tabla(doc,
    ["Situación", "Qué le dices a Claude.ai", "Output esperado"],
    [
        ["Proyecto nuevo",
         '"Necesito automatizar la descarga de X desde el portal Y. '
         'Ya tengo scripts similares para WMS. ¿Cómo estructuro el nuevo?"',
         "Estructura del script: funciones, flujo, manejo de errores, convenciones"],
        ["Decisión de arquitectura",
         '"¿Conviene un agente separado para devoluciones o lo integro en wms_descarga.py?"',
         "Análisis de pros/contras con recomendación concreta"],
        ["Revisar lógica de negocio",
         '"El cálculo de % bloqueado debe excluir pallets en tránsito. ¿Cómo lo modelo en DAX?"',
         "Fórmula DAX con explicación de contexto de filtro"],
        ["Preparar una reunión",
         '"Necesito explicarle a gerencia el impacto del sistema de automatización. '
         'Ayúdame a estructurar los puntos clave."',
         "Estructura de presentación con argumentos y datos"],
        ["Diagnosticar un problema complejo",
         '"El script de staging falla solo los lunes. Estos son los logs: [pega logs]"',
         "Hipótesis de causa raíz y pasos de diagnóstico"],
    ],
    col_widths=[4, 7, 8]
)
doc.add_paragraph()

add_heading(doc, "El puente Claude.ai → Claude Code", 2, color=(68,114,196))
doc.add_paragraph(
    "Claude.ai no guarda tu proyecto. Cuando terminas el diseño en Claude.ai "
    "y pasas a Claude Code para implementar, el puente eres TÚ: copias el plan "
    "o la estructura y se lo das a Claude Code para que lo ejecute."
)
for paso in [
    ("En Claude.ai:", "Diseñas la estructura del script de devoluciones. Claude.ai te da el esquema completo."),
    ("Copias el plan:", "Tomas la estructura que Claude.ai propuso."),
    ("En Claude Code:", '"Implementa este script siguiendo esta estructura: [pegas el plan]". '
                        'Claude Code llama a generador.py y lo construye.'),
    ("Resultado:", "Script completo, con convenciones Egakat, listo en minutos."),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{paso[0]}  ")
    r1.bold = True; r1.font.color.rgb = RGBColor(31, 73, 125)
    p.add_run(paso[1])

doc.add_paragraph()
add_caja(doc, "Importante:",
    "Claude.ai NO tiene acceso a C:\\ClaudeWork\\. Si le muestras código, debes pegarlo tú. "
    "Por eso el filtro de seguridad (ask_ai y los agentes) sanitiza antes de enviar — "
    "pero con Claude.ai en el navegador, TÚ eres el filtro. Nunca pegues el contenido "
    "de tu .env directamente en Claude.ai.")
doc.add_paragraph()
doc.add_page_break()

# ── 3. CÓMO INTERACTÚA CLAUDE CODE CON LOS AGENTES ───────────────────────────
add_heading(doc, "3. Cómo Claude Code orquesta los agentes", 1)
doc.add_paragraph(
    "Claude Code tiene acceso completo al proyecto en C:\\ClaudeWork\\. Cuando le describes "
    "una tarea, internamente decide qué agentes necesita invocar, los llama en el orden correcto "
    "y te entrega el resultado consolidado. Todo esto ocurre automáticamente."
)
doc.add_paragraph()

add_heading(doc, "Ejemplo 1 — Analiza el stock de esta semana", 2, color=(68,114,196))
for paso in [
    ("Tú dices:", '"Analiza el stock WMS de esta semana y dime qué está bloqueado"'),
    ("Claude Code llama:", "extractor.py analizar 'Stock WMS...' → obtiene datos de 23.172 filas"),
    ("Claude Code llama:", "analista.py stock → calcula % bloqueado, top subrubros, anomalías"),
    ("Claude Code llama:", "m365.py subir → guarda el informe en OneDrive"),
    ("Tú recibes:", "Análisis completo con KPIs, alertas y acciones recomendadas"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{paso[0]}  ")
    r1.bold = True; r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(paso[1])
doc.add_paragraph()

add_heading(doc, "Ejemplo 2 — Necesito un script para descargar devoluciones", 2, color=(68,114,196))
for paso in [
    ("Tú dices:", '"Crea un script que descargue el reporte de devoluciones del WMS y lo guarde en OneDrive"'),
    ("Claude Code llama:", "generador.py nuevo 'descargar devoluciones WMS a OneDrive'"),
    ("Claude Code llama:", "generador.py revisar 'devoluciones_descarga.py' → verifica convenciones"),
    ("Claude Code llama:", "generador.py tarea → genera XML para Task Scheduler"),
    ("Tú recibes:", "Script completo + XML de tarea programada listos para usar"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{paso[0]}  ")
    r1.bold = True; r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(paso[1])
doc.add_paragraph()

add_heading(doc, "Ejemplo 3 — Dashboard Power BI de NPS", 2, color=(68,114,196))
for paso in [
    ("Tú dices:", '"Prepara los KPIs DAX para el dashboard de NPS"'),
    ("Claude Code llama:", "m365.py listar 'nps' → busca el archivo más reciente"),
    ("Claude Code llama:", "extractor.py analizar → extrae estructura de datos NPS"),
    ("Claude Code llama:", "power_bi.py kpis nps → genera todas las medidas DAX"),
    ("Claude Code llama:", "power_bi.py modelo → diseña el modelo de datos"),
    ("Tú recibes:", "Medidas DAX listas para pegar en Power BI + modelo de datos"),
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r1 = p.add_run(f"{paso[0]}  ")
    r1.bold = True; r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(paso[1])
doc.add_paragraph()
doc.add_page_break()

# ── 3. CÓMO INTERACTÚA ask_ai CON CLAUDE CODE ────────────────────────────────
add_heading(doc, "4. ask_ai vs Claude Code — cuándo usar cada uno", 1)
doc.add_paragraph(
    "ask_ai y Claude Code NO son lo mismo. Son herramientas complementarias para momentos distintos:"
)
doc.add_paragraph()
add_tabla(doc,
    ["Situación", "Herramienta", "Por qué"],
    [
        ["Estás en terminal y quieres una respuesta rápida", "ask_ai", "No requiere abrir VS Code. Respuesta en segundos."],
        ["Quieres explicar un archivo específico", "ask_ai /explain archivo.py", "Lee el archivo, filtra secrets, pregunta a Claude."],
        ["Trabajas en un proyecto complejo con varios archivos", "Claude Code en VS Code", "Ve todo el proyecto, mantiene contexto entre mensajes."],
        ["Necesitas que un agente haga algo automáticamente", "Claude Code", "Claude Code llama a los agentes por ti."],
        ["Debug rápido de una función puntual", "ask_ai /fix archivo.py", "Diagnóstico en un comando."],
        ["Implementar feature nueva que afecta varios archivos", "Claude Code", "Necesita contexto del proyecto completo."],
        ["Generar tests rápidamente", "ask_ai /test archivo.py", "Genera tests en segundos desde terminal."],
        ["Diseñar arquitectura o analizar lógica", "Claude.ai (navegador)", "Conversación sin límite, sin código real."],
    ],
    col_widths=[5.5, 4.5, 9]
)
doc.add_paragraph()
add_caja(doc, "Resumen:",
    "ask_ai = respuesta rápida desde terminal.  "
    "Claude Code = trabajo profundo con contexto del proyecto.  "
    "Claude.ai = diseño y planificación sin código.")
doc.add_paragraph()
doc.add_page_break()

# ── 4. TU ROL COMO USUARIO ────────────────────────────────────────────────────
add_heading(doc, "5. Tu rol como usuario — qué debes hacer tú", 1)
doc.add_paragraph(
    "El sistema está diseñado para que tú hagas lo mínimo posible. Tu rol es describir, "
    "revisar y aprobar. Los agentes y Claude Code hacen el trabajo técnico."
)
doc.add_paragraph()
add_tabla(doc,
    ["Qué haces tú", "Qué hace el sistema automáticamente"],
    [
        ["Describes el requerimiento en español", "Claude Code interpreta y planifica los pasos"],
        ["Apruebas o rechazas el plan", "Los agentes se ejecutan en el orden correcto"],
        ["Revisas el resultado final", "extractor obtiene datos, analista calcula KPIs, m365 guarda"],
        ["Pides ajustes si es necesario", "generador crea o mejora el script automáticamente"],
        ["Registras la tarea en Task Scheduler", "generador produce el XML listo para registrar"],
        ["Abres Power BI y pegas el DAX", "power_bi genera todas las medidas y el modelo"],
    ],
    col_widths=[8, 11]
)
doc.add_paragraph()
doc.add_page_break()

# ── 5. FLUJO ANTE UN PROYECTO NUEVO ──────────────────────────────────────────
add_heading(doc, "6. Cómo iniciar ante un proyecto o tarea nueva", 1)

add_heading(doc, "Fase 1 — Diseño (Claude.ai en navegador)", 2, color=(68,114,196))
doc.add_paragraph("Antes de escribir código, describe el requerimiento en Claude.ai:")
add_codigo(doc, '"Necesito automatizar X. Tengo acceso a Y. El resultado debe ir a Z. ¿Cómo lo estructuro?"')
doc.add_paragraph("Claude.ai define estructura, funciones y flujo. No ejecuta — solo diseña.")
doc.add_paragraph()

add_heading(doc, "Fase 2 — Implementación (Claude Code en VS Code)", 2, color=(68,114,196))
doc.add_paragraph(
    "Abres VS Code con Claude Code activo. Describes la tarea y Claude Code "
    "coordina los agentes automáticamente:"
)
for item in [
    "Si necesita datos → llama a extractor.py",
    "Si necesita crear un script → llama a generador.py",
    "Si necesita guardar o notificar → llama a m365.py",
    "Si necesita analizar → llama a analista.py",
    "Si necesita DAX o Power BI → llama a power_bi.py",
]:
    doc.add_paragraph(item, style="List Bullet")
doc.add_paragraph()

add_heading(doc, "Fase 3 — Validación rápida (ask_ai desde terminal)", 2, color=(68,114,196))
doc.add_paragraph("Una vez generado el script, valídalo desde terminal sin abrir sesión completa:")
add_codigo(doc, "ask_ai /revisar nuevo_script.py")
add_codigo(doc, "ask_ai /test nuevo_script.py")
doc.add_paragraph()

add_heading(doc, "Fase 4 — Producción", 2, color=(68,114,196))
doc.add_paragraph("Claude Code genera el XML y tú registras la tarea:")
add_codigo(doc, 'py AI_Agent/agentes/generador.py tarea "nuevo_script.py" --hora "08:00"')
add_codigo(doc, 'schtasks /create /tn "Egakat - Nombre" /xml "tarea.xml" /ru "Socrates Cabral" /rp /f')
doc.add_paragraph()
doc.add_page_break()

# ── 6. FLUJO AUTOMÁTICO vs MANUAL ────────────────────────────────────────────
add_heading(doc, "7. Qué es automático y qué requiere tu intervención", 1)
add_tabla(doc,
    ["Acción", "Automático", "Requiere tu intervención"],
    [
        ["Llamar a los agentes correctos", "✅ Claude Code decide", "Solo si quieres cambiar el orden"],
        ["Filtrar datos sensibles antes de enviar a Claude", "✅ Siempre activo", "Nunca"],
        ["Guardar resultados en OneDrive", "✅ m365.py lo hace", "Solo aprobar destino"],
        ["Disparar Power Automate", "✅ Por archivo en carpeta vigilada", "Nunca"],
        ["Registrar tarea en Task Scheduler", "❌", "Tú ejecutas el schtasks"],
        ["Aprobar un script generado antes de producción", "❌", "Siempre — nunca auto-deploy"],
        ["Pegar DAX en Power BI", "❌", "Tú pegas en el editor de medidas"],
        ["Commit y push a GitHub", "❌", "Tú lo haces al cerrar sesión"],
    ],
    col_widths=[6, 5, 8]
)
doc.add_paragraph()
doc.add_page_break()

# ── 7. COMANDOS DE TODOS LOS AGENTES ─────────────────────────────────────────
add_heading(doc, "8. Referencia rápida — todos los comandos", 1)

add_heading(doc, "ask_ai — Consultas rápidas", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ["ask_ai /explain archivo.py", "Explica el código en español"],
        ["ask_ai /fix archivo.py", "Detecta bugs y propone corrección"],
        ["ask_ai /test archivo.py", "Genera tests pytest"],
        ["ask_ai /refactor archivo.py", "Sugiere mejoras de estructura"],
        ['ask_ai "pregunta libre"', "Consulta sin archivo"],
        ['ask_ai "pregunta" archivo.py', "Pregunta específica sobre un archivo"],
    ],
    col_widths=[8, 11]
)
doc.add_paragraph()

add_heading(doc, "extractor — Datos desde cualquier fuente", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ['extractor.py onedrive "Stock WMS"', "Busca archivos en OneDrive"],
        ['extractor.py analizar "archivo.xlsx"', "Extrae + envía a Claude para análisis"],
        ['extractor.py excel "archivo.xlsx"', "Lee Excel y muestra resumen"],
        ['extractor.py pdf "informe.pdf"', "Extrae texto de PDF"],
        ['extractor.py sql "SELECT..." --db "base.db"', "Ejecuta query SQL"],
    ],
    col_widths=[9, 10]
)
doc.add_paragraph()

add_heading(doc, "m365 — OneDrive y Power Automate", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ["m365.py estado", "Estado de todas las carpetas OneDrive"],
        ['m365.py listar "staging" --dias 3', "Archivos recientes en carpeta"],
        ['m365.py subir "reporte.xlsx" --destino "vdr"', "Sube archivo → dispara Power Automate"],
        ['m365.py notificar "mensaje" --flujo alerta', "Escribe trigger → Power Automate actúa"],
        ['m365.py limpiar "staging" --dias 30', "Elimina archivos más antiguos de N días"],
    ],
    col_widths=[9, 10]
)
doc.add_paragraph()

add_heading(doc, "analista — KPIs y análisis logístico", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ['analista.py stock "archivo.xlsx"', "KPIs de inventario: bloqueados, subrubros, alertas"],
        ['analista.py staging "archivo.csv"', "Flujo pallets por cliente"],
        ['analista.py nps "archivo.xlsx"', "Score NPS, promotores, detractores"],
        ['analista.py comparar "a.xlsx" "b.xlsx"', "Diferencias entre dos reportes"],
        ['analista.py informe "stock WMS" --guardar', "Informe completo desde OneDrive + guarda"],
    ],
    col_widths=[9, 10]
)
doc.add_paragraph()

add_heading(doc, "generador — Scripts Python desde requerimiento", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ['generador.py nuevo "descripción"', "Genera script completo con convenciones Egakat"],
        ['generador.py mejorar "script.py" "instrucción"', "Mejora script existente (crea backup)"],
        ['generador.py documentar "script.py"', "Agrega docstrings y type hints"],
        ['generador.py revisar "script.py"', "Auditoría: bugs, seguridad, convenciones"],
        ['generador.py tarea "script.py" --hora "08:00"', "Genera XML para Task Scheduler"],
    ],
    col_widths=[9, 10]
)
doc.add_paragraph()

add_heading(doc, "power_bi — DAX, Power Query y modelos", 2, color=(68,114,196))
add_tabla(doc,
    ["Comando", "Qué hace"],
    [
        ['power_bi.py kpis stock|staging|nps|vdr', "Set completo de medidas DAX por área"],
        ['power_bi.py dax "descripción de medida"', "Genera medidas DAX desde descripción"],
        ['power_bi.py query "transformación a aplicar"', "Genera código Power Query M"],
        ['power_bi.py modelo "fuentes de datos"', "Diseña modelo estrella con relaciones y DAX"],
        ['power_bi.py informe "descripción dashboard"', "Estructura páginas + visuales + medidas"],
    ],
    col_widths=[9, 10]
)
doc.add_paragraph()
doc.add_page_break()

# ── 8. AHORRO DE TOKENS ───────────────────────────────────────────────────────
add_heading(doc, "9. Ahorro de tokens — tus archivos reales", 1)
add_tabla(doc,
    ["Archivo", "Tamaño", "Tokens completo", "Tokens agente", "Ahorro"],
    [
        ["nps_descarga.py",        "18.734 c", "4.684", "1.000", "79%"],
        ["vdr_comparador.py",      "16.307 c", "4.077", "1.000", "75%"],
        ["posiciones_descarga.py", " 9.373 c", "2.343", "1.000", "57%"],
        ["staging_descarga.py",    " 8.069 c", "2.017", "1.000", "50%"],
        ["wms_descarga.py",        " 5.983 c", "1.496", "1.000", "33%"],
        ["TOTAL 7 archivos",       "62.286 c", "15.572", "6.955", "55%"],
    ],
    col_widths=[5, 3, 3.5, 3.5, 2]
)
doc.add_paragraph()
add_caja(doc, "Nota:",
    "El ahorro real en sesiones largas es 2-4x mayor porque el historial se acumula "
    "en cada mensaje. Los archivos pequeños (bajo 400 chars) no se benefician — el agente "
    "agrega overhead de system prompt.")
doc.add_paragraph()
doc.add_page_break()

# ── 9. SEGURIDAD ──────────────────────────────────────────────────────────────
add_heading(doc, "10. Seguridad — filtro automático de credenciales", 1)
doc.add_paragraph(
    "Todos los agentes aplican sanitización antes de enviar a Claude. "
    "El filtro redacta el VALOR pero mantiene el nombre de la variable — "
    "el código sigue siendo legible."
)
add_tabla(doc,
    ["Patrón detectado", "Original", "Lo que recibe Claude"],
    [
        ["password / passwd",  "WMS_PASSWORD=abc123",           "WMS_PASSWORD=[REDACTED]"],
        ["secret",             "client_secret=xyz789",          "client_secret=[REDACTED]"],
        ["api_key / token",    "ANTHROPIC_API_KEY=sk-ant-...", "ANTHROPIC_API_KEY=[REDACTED]"],
        ["private_key",        "PRIVATE_KEY=-----BEGIN...",     "PRIVATE_KEY=[REDACTED]"],
    ],
    col_widths=[4.5, 5.5, 7.5]
)
doc.add_paragraph()
doc.add_page_break()

# ── 10. ESTADO DEL LABORATORIO ────────────────────────────────────────────────
add_heading(doc, "11. Estado actual del laboratorio", 1)
add_tabla(doc,
    ["Componente", "Archivo", "Estado", "Fecha"],
    [
        ["VS Code + Claude Code",     "—",                            "✅ Activo",  "11/03/2026"],
        ["GitHub repositorio",        "socrates-cabral/ClaudeWork-",  "✅ Activo",  "11/03/2026"],
        ["Agente base CLI",           "AI_Agent/ask_ai.py",           "✅ v1.0",    "11/03/2026"],
        ["Agente Extractor",          "AI_Agent/agentes/extractor.py","✅ v1.0",    "11/03/2026"],
        ["Agente M365",               "AI_Agent/agentes/m365.py",     "✅ v1.0",    "11/03/2026"],
        ["Agente Analista",           "AI_Agent/agentes/analista.py", "✅ v1.0",    "11/03/2026"],
        ["Agente Generador",          "AI_Agent/agentes/generador.py","✅ v1.0",    "11/03/2026"],
        ["Agente Power BI",           "AI_Agent/agentes/power_bi.py", "✅ v1.0",    "11/03/2026"],
        ["Dashboard NPS Power BI",    "—",                            "🔲 Pendiente","28/03/2026"],
        ["Azure AD App Registration", "—",                            "🔲 Pendiente IT","—"],
    ],
    col_widths=[5, 6, 3.5, 3.5]
)

# ── PIE ───────────────────────────────────────────────────────────────────────
pie_pagina(doc)
doc.save(SALIDA)
print(f"Documento generado: {SALIDA}")
