import sys
sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# Estilos base
style_normal = doc.styles['Normal']
style_normal.font.name = 'Calibri'
style_normal.font.size = Pt(11)

def titulo(doc, texto, nivel=1):
    h = doc.add_heading(texto, level=nivel)
    h.runs[0].font.color.rgb = RGBColor(0x1A, 0x56, 0x8C)
    return h

def parrafo(doc, texto):
    p = doc.add_paragraph(texto)
    p.paragraph_format.space_after = Pt(6)
    return p

def tabla_simple(doc, headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        c.paragraphs[0].runs[0].bold = True
        c.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1A568C')
        c._tc.get_or_add_tcPr().append(shd)
    # Rows
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            t.rows[i+1].cells[j].text = str(val)
    doc.add_paragraph()

# ─── PORTADA ───
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CRYPTO BOT — GRID TRADING")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8C)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run("Manual de usuario y referencia técnica").font.size = Pt(14)

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.add_run("Socrates Cabral | Abril 2026 | Versión 1.0 — Paper Trading").font.size = Pt(11)

doc.add_page_break()

# ─── 1. QUE ES ESTE BOT ───
titulo(doc, "1. ¿Qué es este bot?")
parrafo(doc, "Es un programa que compra y vende Bitcoin automáticamente, sin que tengas que estar mirando la pantalla. Funciona las 24 horas del día, 7 días a la semana, mientras el computador esté encendido.")
parrafo(doc, "Actualmente está en modo PAPER TRADING — esto significa que simula todas las operaciones con dinero ficticio ($1,000 USDT virtuales). No mueve ni un centavo real hasta que tú lo decidas.")

titulo(doc, "2. ¿Cómo gana dinero?", nivel=1)
parrafo(doc, "Usa una estrategia llamada Grid Trading (trading en cuadrícula). La idea es simple:")
doc.add_paragraph("Divide un rango de precios en niveles equidistantes (como una escalera).", style='List Bullet')
doc.add_paragraph("Cuando el precio BAJA y cruza un nivel → el bot COMPRA Bitcoin en ese nivel.", style='List Bullet')
doc.add_paragraph("Cuando el precio SUBE y cruza ese mismo nivel → el bot VENDE el Bitcoin que compró.", style='List Bullet')
doc.add_paragraph("La diferencia entre el precio de compra y venta es la ganancia.", style='List Bullet')
parrafo(doc, "Ejemplo concreto con la configuración actual:")
tabla_simple(doc,
    ["Evento", "Precio BTC", "Acción", "Resultado"],
    [
        ["BTC baja", "$77,000", "Compra 0.00065 BTC ($50)", "Posición abierta"],
        ["BTC sube", "$78,000", "Vende 0.00065 BTC ($51)", "Ganancia: ~$1 USDT"],
        ["BTC baja otra vez", "$77,000", "Compra de nuevo", "Ciclo se repite"],
    ]
)
parrafo(doc, "Con 20 niveles activos simultáneamente, cada movimiento del precio puede generar múltiples operaciones pequeñas que se acumulan.")

# ─── 3. CONFIGURACION ACTUAL ───
titulo(doc, "3. Configuración actual")
tabla_simple(doc,
    ["Parámetro", "Valor", "¿Qué significa?"],
    [
        ["Par de trading", "BTC/USDT", "Compra y vende Bitcoin contra dólares digitales"],
        ["Rango del grid", "$65,000 – $85,000", "El bot solo opera dentro de este rango de precio"],
        ["Cantidad de niveles", "20", "20 escalones de $1,000 cada uno"],
        ["Capital total", "$1,000 USDT", "Dinero virtual repartido entre los 20 niveles"],
        ["Capital por nivel", "$50 USDT", "$1,000 ÷ 20 = $50 por posición"],
        ["Frecuencia", "Cada 5 minutos", "El bot revisa el mercado cada 5 minutos"],
        ["Exchange", "Crypto.com", "Plataforma de trading conectada"],
        ["Modo", "PAPER TRADING", "Simulación — sin dinero real"],
    ]
)

# ─── 4. EL FILTRO EMA 200 ───
titulo(doc, "4. El filtro EMA 200 (protección de tendencia)")
parrafo(doc, "EMA significa Exponential Moving Average (Media Móvil Exponencial). La EMA 200 calcula el precio promedio de Bitcoin durante los últimos 200 días.")
parrafo(doc, "Su función en el bot:")
doc.add_paragraph("Si BTC está POR ENCIMA de la EMA 200 → mercado alcista → el bot puede COMPRAR y VENDER normalmente.", style='List Bullet')
doc.add_paragraph("Si BTC está POR DEBAJO de la EMA 200 → mercado bajista → el bot solo hace VENTAS, no abre nuevas compras.", style='List Bullet')
parrafo(doc, "¿Por qué? Para no comprar en una tendencia bajista fuerte y quedarse atrapado con posiciones perdedoras.")
parrafo(doc, "Estado actual (Abril 2026):")
tabla_simple(doc,
    ["Indicador", "Valor"],
    [
        ["Precio BTC", "~$77,390"],
        ["EMA 200 diaria", "~$84,107"],
        ["Estado EMA", "BTC bajo EMA → mercado bajista"],
        ["Filtro en este momento", "DESACTIVADO (paper trading — para probar la mecánica)"],
    ]
)

# ─── 5. PROTECCIONES ───
titulo(doc, "5. Protecciones y controles de riesgo")

titulo(doc, "5.1 Drawdown máximo", nivel=2)
parrafo(doc, "Si el bot pierde más del 10% del capital ($100 USDT virtuales), se detiene automáticamente y manda alerta por Telegram. Así no sigue perdiendo.")

titulo(doc, "5.2 Kill Switch", nivel=2)
parrafo(doc, "Si quieres parar el bot inmediatamente, crea el archivo:")
p = doc.add_paragraph()
p.add_run("C:\\ClaudeWork\\crypto_bot\\kill_switch.txt").bold = True
parrafo(doc, "En el próximo ciclo (máximo 5 minutos) el bot detecta este archivo y se detiene de forma limpia. Para reactivarlo, borra ese archivo.")

titulo(doc, "5.3 Máximo de posiciones abiertas", nivel=2)
parrafo(doc, "El bot no puede tener más de 15 posiciones abiertas al mismo tiempo, para no exponer todo el capital de una sola vez.")

titulo(doc, "5.4 Paper Trading obligatorio", nivel=2)
parrafo(doc, "El código tiene un candado: MODO_PAPER_TRADING = True. Cambiarlo a False requiere editar el archivo config.py manualmente. No hay riesgo de activar dinero real por accidente.")

# ─── 6. NOTIFICACIONES ───
titulo(doc, "6. Notificaciones por Telegram")
parrafo(doc, "El bot envía mensajes automáticos a tu Telegram en estos casos:")
tabla_simple(doc,
    ["Evento", "Mensaje que recibes"],
    [
        ["Bot iniciado", "Rango, niveles, capital total"],
        ["Orden ejecutada (BUY)", "Par, precio, cantidad BTC, PnL acumulado"],
        ["Orden ejecutada (SELL)", "Par, precio, cantidad BTC, ganancia de la operación"],
        ["Alerta riesgo", "Tipo de alerta + detalle (drawdown, kill switch)"],
        ["BTC bajo EMA 200", "Aviso de que solo se permiten ventas"],
    ]
)
parrafo(doc, "Todos los mensajes en paper trading llevan el prefijo [PAPER] para que sepas que son simulados.")

# ─── 7. ARCHIVOS IMPORTANTES ───
titulo(doc, "7. Archivos importantes")
tabla_simple(doc,
    ["Archivo", "¿Para qué sirve?"],
    [
        ["config.py", "Todos los parámetros del bot (rango, niveles, capital, etc.)"],
        ["estado_grid.json", "El 'cerebro' del bot — guarda todas las posiciones abiertas y el PnL"],
        ["data/historico_operaciones.json", "Registro completo de todas las órdenes ejecutadas"],
        ["kill_switch.txt", "Si existe, el bot para. Si no existe, el bot corre."],
        ["logs/crypto_bot_*.log", "Registro detallado de cada ciclo (errores, órdenes, precios)"],
        ["run_bot.py", "El programa principal — no necesitas tocarlo"],
    ]
)

# ─── 8. GLOSARIO ───
titulo(doc, "8. Glosario de términos")
tabla_simple(doc,
    ["Término", "Explicación simple"],
    [
        ["Grid Trading", "Estrategia de comprar bajo y vender alto en múltiples niveles fijos"],
        ["USDT", "Dólar digital (stablecoin) — siempre vale ~$1 USD"],
        ["BTC", "Bitcoin"],
        ["EMA 200", "Precio promedio de los últimos 200 días — indica si el mercado sube o baja"],
        ["Paper Trading", "Simulación con dinero ficticio — sin riesgo real"],
        ["PnL", "Profit and Loss — ganancia o pérdida acumulada"],
        ["Drawdown", "Cuánto ha bajado el capital desde su punto más alto"],
        ["Kill Switch", "Interruptor de emergencia para parar el bot"],
        ["Exchange", "Plataforma de intercambio (ej: Crypto.com, Kraken)"],
        ["Nivel / Step", "Cada escalón del grid — en este caso cada $1,000"],
        ["grid_activo", "True = bot puede comprar y vender. False = solo puede vender"],
    ]
)

# ─── 9. FASES ───
titulo(doc, "9. Fases del proyecto")
tabla_simple(doc,
    ["Fase", "Cuándo", "Condición para avanzar"],
    [
        ["Paper Trading (actual)", "Desde Abril 2026", "Mínimo 30 días de operación"],
        ["Capital real pequeño", "+30 días", "PnL paper positivo, sin errores técnicos"],
        ["Escala", "+60 días", "Consistencia confirmada, ROI > 0 sostenido"],
        ["VPS (servidor 24/7)", "Pendiente", "Cuando el laptop sea limitante — Hetzner CX32 ~€7/mes"],
    ]
)

# ─── 10. COMO REVISAR ───
titulo(doc, "10. ¿Cómo saber si el bot está funcionando?")
doc.add_paragraph("Opción 1 — Programador de Tareas de Windows: busca 'Crypto Bot - Grid Trading' → Estado debe ser 'Listo' y 'Última ejecución: hace menos de 5 min'", style='List Number')
doc.add_paragraph("Opción 2 — Logs: abre C:\\ClaudeWork\\logs\\ y busca el archivo crypto_bot más reciente. Debe tener entradas cada 5 minutos.", style='List Number')
doc.add_paragraph("Opción 3 — Telegram: cuando el bot ejecute una orden, recibirás un mensaje automático.", style='List Number')
doc.add_paragraph("Opción 4 — estado_grid.json: ábrelo con Notepad. El campo 'ultima_actualizacion' muestra cuándo corrió por última vez.", style='List Number')

doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("— Documento generado automáticamente por Claude Code —\nAbril 2026").font.size = Pt(9)

output = Path("crypto_bot/Manual_CryptoBot_GridTrading.docx")
doc.save(output)
print(f"Guardado: {output.resolve()}")
