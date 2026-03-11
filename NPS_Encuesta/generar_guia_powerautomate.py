"""
Genera: Guia_PowerAutomate_NPS_Alerta.docx
Instrucciones paso a paso para crear el flow de alerta NPS en Power Automate.
Egakat SPA — Control de Gestión y Mejora Continua
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.5)
    s.left_margin = s.right_margin = Cm(3.0)

AZUL   = RGBColor(0x1F, 0x49, 0x7D)
AZUL2  = RGBColor(0x2E, 0x75, 0xB6)
VERDE  = RGBColor(0x37, 0x86, 0x44)
NARAN  = RGBColor(0xED, 0x7D, 0x31)
GRIS   = RGBColor(0x40, 0x40, 0x40)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)
ROJO   = RGBColor(0xC0, 0x00, 0x00)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def set_borders(cell, color="BFBFBF", sz="4"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcB.append(el)
    tcPr.append(tcB)

def h1(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot); pPr.append(pBdr)

def h2(txt, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(12)
    r.font.color.rgb = color or AZUL2

def body(txt, size=10.5, space=4, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt)
    r.font.size = Pt(size)
    r.font.color.rgb = GRIS
    r.italic = italic

def paso(numero, titulo, descripcion, valor_campo=None, nota=None):
    """Bloque visual para cada paso del flow."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_bg(cell, "F0F7FF")
    set_borders(cell, color="2E75B6", sz="6")

    # Numero + titulo
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f"Paso {numero} — ")
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = AZUL
    r2 = p.add_run(titulo)
    r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = AZUL

    # Descripcion
    p2 = cell.add_paragraph()
    p2.paragraph_format.left_indent = Cm(0.3)
    p2.paragraph_format.space_after = Pt(2)
    r3 = p2.add_run(descripcion)
    r3.font.size = Pt(10.5); r3.font.color.rgb = GRIS

    # Valor del campo (en caja gris)
    if valor_campo:
        p3 = cell.add_paragraph()
        p3.paragraph_format.left_indent = Cm(0.5)
        p3.paragraph_format.space_after = Pt(4)
        r4 = p3.add_run(valor_campo)
        r4.font.size = Pt(10)
        r4.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
        r4.bold = True

    # Nota opcional
    if nota:
        p4 = cell.add_paragraph()
        p4.paragraph_format.left_indent = Cm(0.3)
        p4.paragraph_format.space_after = Pt(4)
        r5 = p4.add_run(f"Nota: {nota}")
        r5.font.size = Pt(9.5)
        r5.italic = True
        r5.font.color.rgb = NARAN

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def campo(nombre, valor, obligatorio=True):
    """Fila de tabla campo/valor."""
    return [
        ("* " if obligatorio else "  ") + nombre,
        valor
    ]

def tabla_campos(filas):
    tbl = doc.add_table(rows=1+len(filas), cols=2)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    for i, h in enumerate(["Campo", "Valor a ingresar"]):
        c = tbl.rows[0].cells[i]
        set_bg(c, "1F497D")
        p = c.paragraphs[0]
        r = p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=BLANCO

    for ri, (campo_txt, valor_txt) in enumerate(filas):
        bg = "EBF3FB" if ri%2==0 else "FFFFFF"
        c0 = tbl.rows[ri+1].cells[0]; c1 = tbl.rows[ri+1].cells[1]
        set_bg(c0, bg); set_bg(c1, bg)
        r0 = c0.paragraphs[0].add_run(campo_txt)
        r0.font.size=Pt(10); r0.font.color.rgb=GRIS
        if campo_txt.startswith("*"):
            r0.bold=True
        r1 = c1.paragraphs[0].add_run(valor_txt)
        r1.font.size=Pt(10); r1.font.color.rgb=RGBColor(0x1F,0x49,0x7D)
        r1.bold=True

    tbl.columns[0].width = Cm(5)
    tbl.columns[1].width = Cm(10)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def callout(txt, bg="EAF4FB", border="2E75B6"):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0,0)
    set_bg(cell, bg)
    set_borders(cell, color=border, sz="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(txt); r.font.size=Pt(10.5); r.font.color.rgb=GRIS
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# =============================================================================
# PORTADA
# =============================================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EGAKAT SPA"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=AZUL2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Guia de Configuracion")
r.bold=True; r.font.size=Pt(22); r.font.color.rgb=AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Flow Power Automate — Alerta NPS Sin Respuestas")
r.bold=True; r.font.size=Pt(15); r.font.color.rgb=AZUL2

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Paso a paso para crear el flujo de notificacion automatica")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRIS

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_table(rows=4, cols=2)
info.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(lbl,val) in enumerate([
    ("Preparado por:", "Control de Gestion y Mejora Continua"),
    ("Fecha:",         datetime.date.today().strftime("%d de %B de %Y")),
    ("Plataforma:",    "Power Automate — make.powerautomate.com"),
    ("Version:",       "1.0"),
]):
    c0=info.rows[i].cells[0]; c1=info.rows[i].cells[1]
    set_bg(c0,"1F497D"); set_bg(c1,"EBF3FB" if i%2==0 else "FFFFFF")
    r0=c0.paragraphs[0].add_run(lbl); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=BLANCO
    r1=c1.paragraphs[0].add_run(val); r1.font.size=Pt(10); r1.font.color.rgb=GRIS

doc.add_page_break()

# =============================================================================
# RESUMEN DEL FLOW
# =============================================================================
h1("Resumen del flujo")

body(
    "Este flow se activa automaticamente cuando el script nps_descarga.py "
    "detecta que la encuesta no tiene respuestas y crea un archivo de alerta "
    "en OneDrive. Power Automate lee ese archivo y envia un correo a Socrates "
    "y Franco con el detalle del problema."
)

# Diagrama de flujo como tabla
tbl_flow = doc.add_table(rows=6, cols=1)
tbl_flow.alignment = WD_TABLE_ALIGNMENT.CENTER
pasos_flow = [
    ("TRIGGER", "Se crea un archivo en OneDrive\n/Reportes NPS/Alertas/", "2E75B6"),
    ("PASO 1",  "Esperar 2 minutos\n(para que el archivo sincronice con la nube)", "70AD47"),
    ("PASO 2",  "Obtener contenido del archivo\n(Get file content)", "70AD47"),
    ("PASO 3",  "Obtener propiedades del archivo\n(Get file metadata — para el nombre)", "70AD47"),
    ("PASO 4",  "Enviar correo electronico\nSocrates + Franco", "70AD47"),
    ("FIN",     "Correo recibido con detalle de la alerta", "1F497D"),
]
for i, (etq, desc, color) in enumerate(pasos_flow):
    c = tbl_flow.rows[i].cells[0]
    set_bg(c, color)
    set_borders(c, color="FFFFFF", sz="6")
    p = c.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f"{etq}  ")
    r1.bold=True; r1.font.size=Pt(10); r1.font.color.rgb=BLANCO
    r2 = p.add_run(desc)
    r2.font.size=Pt(10); r2.font.color.rgb=BLANCO

doc.add_paragraph().paragraph_format.space_after = Pt(6)
doc.add_page_break()

# =============================================================================
# PASOS DETALLADOS
# =============================================================================
h1("Configuracion paso a paso")

body("Ingresa a Power Automate: make.powerautomate.com con tu cuenta egakat.")
doc.add_paragraph()

# CREAR FLOW
h2("Crear el flow")
paso(
    numero=1,
    titulo="Crear nuevo flow automatizado",
    descripcion="En el panel izquierdo haz clic en 'Crear' → selecciona 'Flujo de nube automatizado'.",
    valor_campo="Nombre del flow: NPS Egakat - Alerta Sin Respuestas",
    nota="Si te pide buscar un conector de trigger, escribe 'OneDrive' y selecciona el trigger del Paso 2."
)

# TRIGGER
h2("Trigger — Cuando se crea el archivo de alerta")
paso(
    numero=2,
    titulo="Trigger: When a file is created — OneDrive for Business",
    descripcion="Busca el conector 'OneDrive for Business' y selecciona el trigger:",
    valor_campo="When a file is created (solo archivos, no carpetas)",
)
body("Configura el trigger con estos valores:", size=10.5)
tabla_campos([
    ("* Folder (Carpeta)", "/Reportes NPS/Alertas"),
])
callout(
    "Si la carpeta /Alertas no aparece en el selector, creala primero en OneDrive "
    "o ejecuta el script una vez para que la cree automaticamente.",
    bg="FFF2CC", border="ED7D31"
)

# DELAY
h2("Paso 1 — Esperar sincronizacion")
paso(
    numero=3,
    titulo="Delay — Esperar 2 minutos",
    descripcion="Agrega la accion 'Delay' (busca 'delay' en el buscador de acciones).",
)
tabla_campos([
    ("* Count (Cantidad)", "2"),
    ("* Unit (Unidad)",    "Minute"),
])

# GET FILE CONTENT
h2("Paso 2 — Leer el archivo de alerta")
paso(
    numero=4,
    titulo="Get file content — OneDrive for Business",
    descripcion=(
        "Agrega la accion 'Get file content' de OneDrive for Business. "
        "Esto lee el texto del archivo .txt que creo el script."
    ),
)
tabla_campos([
    ("* File (Archivo)", "Id — seleccionar de contenido dinamico del trigger\n"
                         "Haz clic en el campo → 'Contenido dinamico' → 'Id'"),
    ("  Infer Content Type", "Yes"),
])

# GET FILE METADATA
h2("Paso 3 — Obtener nombre del archivo")
paso(
    numero=5,
    titulo="Get file metadata — OneDrive for Business",
    descripcion=(
        "Agrega la accion 'Get file metadata'. Nos permite mostrar el nombre "
        "del archivo en el correo para identificar la fecha de la alerta."
    ),
)
tabla_campos([
    ("* File (Archivo)", "Id — contenido dinamico del trigger → 'Id'"),
])

# SEND EMAIL
h2("Paso 4 — Enviar correo de alerta")
paso(
    numero=6,
    titulo="Send an email (V2) — Office 365 Outlook",
    descripcion=(
        "Agrega la accion 'Send an email (V2)' del conector Office 365 Outlook. "
        "Configura los campos exactamente como se indica:"
    ),
)
tabla_campos([
    ("* To (Para)",    "socrates.cabral@egakat.cl; franco.perez@egakat.cl"),
    ("* Subject (Asunto)", "Alerta NPS — Sin respuestas al momento de descarga"),
    ("* Body (Cuerpo)", "Ver detalle abajo"),
    ("  Importance",   "High"),
])

body("Cuerpo del correo — copiar exactamente este texto y completar con contenido dinamico:", size=10.5)

# Caja con el cuerpo del correo
callout(
    "Hola,\n\n"
    "El script de descarga NPS se ejecuto y no encontro respuestas en la encuesta.\n\n"
    "Detalle de la alerta:\n"
    "[Insertar aqui: contenido dinamico → 'File Content' del Paso 2]\n\n"
    "Archivo: [Insertar aqui: contenido dinamico → 'Name' del Paso 3]\n\n"
    "Acciones recomendadas:\n"
    "  - Verificar que el envio de la encuesta fue realizado en LimeSurvey.\n"
    "  - Verificar que el periodo de respuesta no haya cerrado sin respuestas.\n"
    "  - Si el envio aun no se realizo, programarlo en LimeSurvey.\n\n"
    "Este correo fue generado automaticamente por el sistema NPS Egakat.\n\n"
    "Egakat SPA — Control de Gestion y Mejora Continua",
    bg="F5F5F5", border="BFBFBF"
)

callout(
    "Como insertar contenido dinamico en el cuerpo:\n"
    "1. Haz clic dentro del campo Body donde dice [Insertar aqui...]\n"
    "2. Aparece el panel 'Contenido dinamico' a la derecha\n"
    "3. Para el contenido del archivo: busca 'File Content' bajo 'Get file content'\n"
    "4. Para el nombre: busca 'Name' bajo 'Get file metadata'\n"
    "5. Haz clic en la variable para insertarla",
    bg="EAF4FB", border="2E75B6"
)

doc.add_page_break()

# =============================================================================
# GUARDAR Y PROBAR
# =============================================================================
h1("Guardar y probar el flow")

paso(
    numero=7,
    titulo="Guardar el flow",
    descripcion="Haz clic en 'Guardar' en la esquina superior derecha. "
                "Power Automate validara que no haya errores en los conectores.",
    nota="Si pide autorizacion para OneDrive o Outlook, haz clic en 'Iniciar sesion' "
         "y autoriza con tu cuenta egakat."
)

paso(
    numero=8,
    titulo="Probar el flow manualmente",
    descripcion=(
        "Para verificar que funciona antes de esperar una alerta real:\n"
        "1. Haz clic en 'Probar' (esquina superior derecha)\n"
        "2. Selecciona 'Manualmente'\n"
        "3. Crea un archivo .txt cualquiera en OneDrive en la carpeta /Reportes NPS/Alertas/\n"
        "4. El flow debe ejecutarse y llegar el correo a ambos destinatarios en ~3 minutos"
    ),
    nota="Puedes crear el archivo de prueba desde el explorador de archivos de Windows "
         "en la carpeta de OneDrive sincronizada."
)

paso(
    numero=9,
    titulo="Activar el flow",
    descripcion=(
        "Si el flow quedo en estado 'Desactivado' despues de guardarlo, "
        "ve a la lista de flows → busca 'NPS Egakat - Alerta Sin Respuestas' "
        "→ haz clic en los tres puntos → 'Activar'."
    )
)

# =============================================================================
# REFERENCIA RAPIDA
# =============================================================================
h1("Referencia rapida — resumen del flow")

tbl_ref = doc.add_table(rows=1, cols=4)
tbl_ref.style = "Table Grid"
for i,h in enumerate(["Paso","Conector","Accion","Campo clave"]):
    c = tbl_ref.rows[0].cells[i]
    set_bg(c, "1F497D")
    r = c.paragraphs[0].add_run(h)
    r.bold=True; r.font.size=Pt(10); r.font.color.rgb=BLANCO

resumen = [
    ["Trigger", "OneDrive for Business",  "When a file is created",   "Folder: /Reportes NPS/Alertas"],
    ["1",       "Schedule",               "Delay",                     "2 minutos"],
    ["2",       "OneDrive for Business",  "Get file content",          "File: Id (dinamico)"],
    ["3",       "OneDrive for Business",  "Get file metadata",         "File: Id (dinamico)"],
    ["4",       "Office 365 Outlook",     "Send an email (V2)",        "To: socrates + franco"],
]
for ri, fila in enumerate(resumen):
    bg = "EBF3FB" if ri%2==0 else "FFFFFF"
    row = tbl_ref.add_row()
    for ci, val in enumerate(fila):
        c = row.cells[ci]
        set_bg(c, bg)
        r = c.paragraphs[0].add_run(val)
        r.font.size=Pt(10); r.font.color.rgb=GRIS

doc.add_paragraph()

# Cierre
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(
    f"Egakat SPA  |  Control de Gestion y Mejora Continua  |  "
    f"{datetime.date.today().strftime('%B %Y')}"
)
r.font.size=Pt(9); r.italic=True; r.font.color.rgb=RGBColor(0x80,0x80,0x80)

salida = r"C:\ClaudeWork\NPS_Encuesta\Guia_PowerAutomate_NPS_Alerta.docx"
doc.save(salida)
print(f"Documento generado: {salida}")
