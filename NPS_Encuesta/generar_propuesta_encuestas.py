"""
Genera: Propuesta_Configuracion_Encuestas_NPS_CSAT.docx
Egakat SPA — Control Management & Continuous Improvement
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# Margenes
for s in doc.sections:
    s.top_margin = s.bottom_margin = Cm(2.5)
    s.left_margin = s.right_margin = Cm(3.0)

# Colores
AZUL   = RGBColor(0x1F, 0x49, 0x7D)
AZUL2  = RGBColor(0x2E, 0x75, 0xB6)
VERDE  = RGBColor(0x37, 0x86, 0x44)
ROJO   = RGBColor(0xC0, 0x00, 0x00)
NARAN  = RGBColor(0xED, 0x7D, 0x31)
GRIS   = RGBColor(0x40, 0x40, 0x40)
BLANCO = RGBColor(0xFF, 0xFF, 0xFF)

def set_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def h1(txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(14); r.font.color.rgb = AZUL
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
    bot.set(qn("w:space"), "4"); bot.set(qn("w:color"), "1F497D")
    pBdr.append(bot); pPr.append(pBdr)

def h2(txt, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(txt)
    r.bold = True; r.font.size = Pt(11.5)
    r.font.color.rgb = color or AZUL2

def body(txt, size=10.5, space=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space)
    r = p.add_run(txt)
    r.font.size = Pt(size); r.font.color.rgb = GRIS

def bullet(txt, color=None, bold=False):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(txt)
    r.font.size = Pt(10.5)
    r.font.color.rgb = color or GRIS
    r.bold = bold

def callout(txt, bg="EAF4FB", border="2E75B6"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_bg(cell, bg)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcB = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),"single"); el.set(qn("w:sz"),"8")
        el.set(qn("w:space"),"0"); el.set(qn("w:color"), border)
        tcB.append(el)
    tcPr.append(tcB)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    r = p.add_run(txt); r.font.size = Pt(10.5); r.font.color.rgb = GRIS
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

def tabla(headers, filas, hdr_bg="1F497D", alt_bg="EBF3FB"):
    tbl = doc.add_table(rows=1+len(filas), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]
        set_bg(c, hdr_bg)
        p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h); r.bold=True; r.font.size=Pt(10); r.font.color.rgb=BLANCO
    for ri, fila in enumerate(filas):
        bg = alt_bg if ri%2==0 else "FFFFFF"
        for ci, val in enumerate(fila):
            c = tbl.rows[ri+1].cells[ci]
            set_bg(c, bg)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci>0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val)); r.font.size=Pt(10); r.font.color.rgb=GRIS
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def tag(txt, color_bg, color_txt="FFFFFF"):
    """Pequeño tag de color inline (tabla 1x1 chica)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    c = tbl.cell(0,0)
    set_bg(c, color_bg)
    p = c.paragraphs[0]
    r = p.add_run(f"  {txt}  ")
    r.bold=True; r.font.size=Pt(9)
    r.font.color.rgb = RGBColor.from_string(color_txt)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# =============================================================================
# PORTADA
# =============================================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("EGAKAT SPA"); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=AZUL2

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Propuesta de Configuracion")
r.bold=True; r.font.size=Pt(22); r.font.color.rgb=AZUL

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Encuestas NPS y CSAT — LimeSurvey")
r.bold=True; r.font.size=Pt(16); r.font.color.rgb=AZUL2

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Comparativa estado actual vs. configuracion propuesta")
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRIS

doc.add_paragraph()
doc.add_paragraph()

info_tbl = doc.add_table(rows=4, cols=2)
info_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
for i,(lbl,val) in enumerate([
    ("Preparado por:", "Control Management & Continuous Improvement"),
    ("Fecha:",         datetime.date.today().strftime("%d de %B de %Y")),
    ("Plataforma:",    "LimeSurvey Cloud v6.2.11 — egakat.limesurvey.net"),
    ("Version:",       "1.0"),
]):
    c0=info_tbl.rows[i].cells[0]; c1=info_tbl.rows[i].cells[1]
    set_bg(c0,"1F497D"); set_bg(c1,"EBF3FB" if i%2==0 else "FFFFFF")
    r0=c0.paragraphs[0].add_run(lbl); r0.bold=True; r0.font.size=Pt(10); r0.font.color.rgb=BLANCO
    r1=c1.paragraphs[0].add_run(val); r1.font.size=Pt(10); r1.font.color.rgb=GRIS

doc.add_page_break()

# =============================================================================
# 1. CONTEXTO
# =============================================================================
h1("1. Contexto y objetivo")
body(
    "Egakat cuenta actualmente con dos encuestas activas en LimeSurvey. "
    "Este documento describe el estado actual de cada una, los problemas "
    "identificados y la configuracion propuesta para maximizar el valor del "
    "instrumento NPS y del diagnostico operacional CSAT."
)
callout(
    "Objetivo: que el NPS sea un indicador gerencial confiable, comparable "
    "con benchmarks de industria 3PL, y que el CSAT operacional entregue "
    "informacion accionable mensualmente para el equipo de operaciones.",
    bg="EAF4FB", border="2E75B6"
)

tabla(
    ["Encuesta", "ID LimeSurvey", "Proposito", "Frecuencia propuesta"],
    [
        ["Encuesta NPS — '¿Nos recomendarias?'", "418429", "KPI de lealtad para gerencia", "Trimestral"],
        ["Encuesta CSAT — 'Satisfaccion mensual'", "386641", "Diagnostico operacional por area", "Mensual"],
    ]
)

# =============================================================================
# 2. ENCUESTA NPS (ID 418429)
# =============================================================================
h1("2. Encuesta NPS — ID 418429")

h2("2.1 Estado actual (lo que Franco configuro)")
body("Preguntas detectadas via API el 10/03/2026:")

tabla(
    ["Codigo", "Tipo", "Pregunta actual", "Observacion"],
    [
        ["Q00",    "Lista radio (L)", "En una escala del 0 al 10, que tan probable es que recomiendes...", "OK — escala ya corregida a 0-10"],
        ["G01Q02", "Multiple (M)",    "Que criterios son importantes para ti al elegir un servicio...\n(8 opciones: Calidad, Atencion, Reputacion, Rapidez, Precision, Flexibilidad, Comunicacion, Experiencia)", "NO es pregunta NPS estandar — es diagnostico de drivers"],
        ["G01Q03", "Texto largo (U)", "Por favor comparte cualquier comentario o sugerencia adicional...", "Util, pero muy generica"],
        ["G01Q04", "Lista radio (L)", "Te gustaria que nos comuniquemos contigo para profundizar tus respuestas?", "Buena — contacto condicional"],
        ["G01Q05", "Texto corto (Q)", "Como prefieres que te contactemos?", "Falta capturar: nombre, empresa, cargo, email"],
    ]
)

h2("2.2 Problemas identificados", color=ROJO)
bullet("G01Q02 (criterios) NO es una pregunta NPS. Diagnostica lo que valoran los clientes, pero no explica por que dieron ese puntaje especifico. Pertenece a CSAT.", color=ROJO)
bullet("G01Q03 (comentarios) es muy abierta. Los clientes no saben si responder sobre el servicio general o sobre su puntuacion.", color=NARAN)
bullet("G01Q05 solo captura 'como contactar' pero no los datos de contacto. Si el cliente dice Si, necesitas nombre, empresa, cargo y email.", color=NARAN)
doc.add_paragraph()

h2("2.3 Configuracion propuesta", color=VERDE)
callout(
    "La encuesta NPS debe tener 3 preguntas. Breve, directa y enfocada en el score. "
    "El objetivo es maximizar la tasa de respuesta y obtener el NPS puro.",
    bg="E2EFDA", border="378644"
)

tabla(
    ["N°", "Codigo", "Tipo LimeSurvey", "Pregunta propuesta", "Cambio requerido"],
    [
        ["1", "Q00",    "Lista radio (L)\n11 opciones: 0 a 10",
         "En una escala de 0 (nada probable) a 10 (extremadamente probable), que tan probable es que recomiendes los servicios de Egakat a un colega o amigo?",
         "Mantener — ya esta correcto"],
        ["2", "G01Q02", "Texto corto (S) — REEMPLAZAR",
         "Cual es la principal razon de tu puntuacion? (opcional)",
         "Eliminar seleccion multiple. Reemplazar por texto libre corto. Marcar como opcional."],
        ["3", "G01Q04", "Si/No (Y)",
         "Te gustaria que nos contactemos contigo para conocer mas sobre tu experiencia?",
         "Cambiar tipo a Si/No. Agregar logica condicional: si Si -> mostrar G01Q05."],
        ["4", "G01Q05", "Multiple short text (Q) — EXPANDIR",
         "Nombre / Empresa / Cargo / Correo electronico",
         "Agregar subcampos: nombre, empresa, cargo, email. Solo visible si G01Q04 = Si."],
    ]
)

body("Que hacer con G01Q02 actual (criterios):")
p = doc.add_paragraph()
r = p.add_run("Mover la pregunta de criterios a la Encuesta CSAT (ID 386641). "
              "Ahi tiene mas sentido como analisis de drivers de satisfaccion.")
r.font.size = Pt(10.5); r.font.color.rgb = GRIS

h2("2.4 Configuracion de logica condicional (relevance equation)")
callout(
    "En LimeSurvey: editar pregunta G01Q05 -> campo 'Relevance equation':\n\n"
    "    {G01Q04} == 'Y'\n\n"
    "Esto hace que los campos de contacto solo aparezcan si el cliente responde Si.",
    bg="FFF2CC", border="ED7D31"
)

h2("2.5 Configuracion de envio — frecuencia trimestral")
tabla(
    ["Parametro", "Configuracion"],
    [
        ["Frecuencia de envio",  "Trimestral (Marzo, Junio, Septiembre, Diciembre)"],
        ["Audiencia",            "Todos los clientes activos — lista de tokens en LimeSurvey"],
        ["Anonimato",            "Activado — mejora tasa de respuesta"],
        ["Tiempo estimado",      "Menos de 1 minuto"],
        ["Recordatorio",         "1 recordatorio a los 7 dias si no respondio"],
        ["Canal",                "Email desde LimeSurvey (configurar remitente egakat)"],
    ]
)

doc.add_page_break()

# =============================================================================
# 3. ENCUESTA CSAT (ID 386641)
# =============================================================================
h1("3. Encuesta CSAT Operacional — ID 386641")

h2("3.1 Estado actual")
body("Preguntas detectadas via API el 10/03/2026:")

tabla(
    ["Codigo", "Tipo", "Pregunta actual", "Estado"],
    [
        ["Q00",    "Lista radio (L)", "En general, que tan satisfecho estas con el servicio logistico de Egakat?", "OK"],
        ["G01Q02", "Array (A)",       "Como calificarias las siguientes areas: Recepcion, Preparacion, Despacho, Inventarios, Transporte, Calidad, Servicio al cliente", "OK — buena estructura"],
        ["G01Q03", "Texto (T)",       "Que aspecto del servicio valoras mas?", "OK"],
        ["G01Q04", "Lista radio (L)", "Experimentaste algun problema con nuestro servicio?", "Falta skip logic"],
        ["G01Q05", "Texto (T)",       "Por favor describenos el problema", "Debe ocultarse si G01Q04 = No"],
        ["G01Q11", "Lista radio (L)", "Con que frecuencia logramos resolver tus problemas satisfactoriamente?", "Debe ocultarse si G01Q04 = No"],
        ["G01Q06", "Escala 1-5",      "Nuestro equipo resolvio tus dudas o problemas oportunamente?", "Debe ocultarse si G01Q04 = No"],
        ["G01Q14", "Texto (T)",       "Que podriamos mejorar?", "OK"],
        ["G01Q15", "Escala 1-5",      "Que tan satisfecho estas con el cumplimiento de tiempos de entrega?", "OK — SLA"],
        ["G01Q16", "Escala 1-5",      "Que tan satisfecho estas con la precision de los pedidos?", "OK"],
        ["G01Q17", "Escala 1-5",      "Que tan satisfecho estas con la informacion y seguimiento?", "OK"],
    ]
)

h2("3.2 Problemas identificados", color=ROJO)
bullet("G01Q05, G01Q11 y G01Q06 se muestran siempre, aunque el cliente diga que no tuvo problemas. Genera friccion y puede provocar abandono.", color=ROJO)
bullet("Falta la pregunta de criterios (actualmente en NPS). Si se mueve aqui, enriquece el diagnostico.", color=NARAN)
bullet("No hay campo de Empresa/Cargo del encuestado — sin esto no se puede segmentar el CSAT por cliente.", color=NARAN)
doc.add_paragraph()

h2("3.3 Configuracion propuesta", color=VERDE)

tabla(
    ["Pregunta", "Cambio requerido", "Prioridad"],
    [
        ["G01Q05 (descripcion problema)", "Agregar relevance: {G01Q04} == 'Y'", "Alta"],
        ["G01Q11 (frecuencia resolucion)", "Agregar relevance: {G01Q04} == 'Y'", "Alta"],
        ["G01Q06 (resolucion oportuna)",   "Agregar relevance: {G01Q04} == 'Y'", "Alta"],
        ["Nueva pregunta — Criterios",     "Mover G01Q02 desde NPS: que criterios son importantes al elegir un servicio logistico?", "Media"],
        ["Nueva pregunta — Empresa",       "Agregar campo texto: Empresa / razon social del encuestado", "Alta"],
        ["Nueva pregunta — Cargo",         "Agregar campo texto: Cargo del encuestado", "Media"],
    ]
)

h2("3.4 Frecuencia de envio — mensual")
callout(
    "La CSAT operacional debe enviarse mensualmente para detectar problemas "
    "a tiempo. Con frecuencia semestral, un problema puede durar 6 meses "
    "sin que aparezca en el radar. Con frecuencia mensual, se detecta y "
    "se corrige en el ciclo siguiente.",
    bg="E2EFDA", border="378644"
)

tabla(
    ["Parametro", "Configuracion"],
    [
        ["Frecuencia de envio",  "Mensual — primeros 5 dias del mes"],
        ["Audiencia",            "Clientes activos con operacion en el mes anterior"],
        ["Anonimato",            "Desactivado — necesitamos saber quien responde para segmentar"],
        ["Tiempo estimado",      "3-5 minutos"],
        ["Recordatorio",         "1 recordatorio a los 5 dias"],
        ["Canal",                "Email desde LimeSurvey"],
    ]
)

doc.add_page_break()

# =============================================================================
# 4. RESUMEN DE CAMBIOS
# =============================================================================
h1("4. Resumen de cambios a solicitar a Franco")

h2("Encuesta NPS (ID 418429) — 4 cambios")
tabla(
    ["#", "Que cambiar", "Como hacerlo en LimeSurvey", "Prioridad"],
    [
        ["1", "Reemplazar G01Q02 (criterios multiple) por texto libre 'Razon de tu puntuacion'",
         "Editar pregunta -> cambiar tipo a Short text (S) -> cambiar texto -> marcar Opcional",
         "Alta"],
        ["2", "Cambiar G01Q04 a tipo Si/No",
         "Editar pregunta -> cambiar tipo a Yes/No (Y)",
         "Alta"],
        ["3", "Agregar datos de contacto en G01Q05 (nombre, empresa, cargo, email)",
         "Editar pregunta -> agregar subpreguntas -> agregar relevance equation: {G01Q04}=='Y'",
         "Alta"],
        ["4", "Mover pregunta de criterios a encuesta CSAT",
         "Copiar pregunta desde NPS -> pegar en CSAT (menu de pregunta -> Copiar)",
         "Media"],
    ]
)

h2("Encuesta CSAT (ID 386641) — 3 cambios")
tabla(
    ["#", "Que cambiar", "Como hacerlo en LimeSurvey", "Prioridad"],
    [
        ["1", "Agregar skip logic a G01Q05, G01Q11, G01Q06",
         "Editar cada pregunta -> campo 'Relevance equation' -> escribir: {G01Q04} == 'Y'",
         "Alta"],
        ["2", "Agregar campos Empresa y Cargo",
         "Agregar nueva pregunta tipo Short text al final del formulario",
         "Alta"],
        ["3", "Pegar pregunta de criterios traida desde NPS",
         "Ver punto 4 de cambios NPS",
         "Media"],
    ]
)

# =============================================================================
# 5. IMPACTO EN POWER BI
# =============================================================================
h1("5. Impacto en Power BI — como se conecta todo")

body(
    "El script nps_descarga.py se conecta a LimeSurvey via API, descarga las "
    "respuestas y genera un Excel en OneDrive con 3 hojas. Power BI lee ese Excel "
    "sin necesidad de Azure AD ni aprobacion de TI."
)

tabla(
    ["Hoja Excel", "Contenido", "Uso en Power BI"],
    [
        ["Resumen NPS",    "Score, % promotores/pasivos/detractores, promedio", "Tarjetas KPI y gauge"],
        ["Respuestas",     "Detalle individual + columna categoria NPS calculada", "Tabla de detalle y filtros"],
        ["PowerBI_Datos",  "Fila unica estructurada por fecha de extraccion", "Grafico de tendencia historica"],
    ]
)

callout(
    "Para conectar Power BI: Obtener datos -> Excel -> seleccionar archivo en OneDrive "
    "-> seleccionar hoja 'PowerBI_Datos'. Configurar actualizacion automatica desde Power BI Service.",
    bg="EAF4FB", border="2E75B6"
)

# Cierre
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run(f"Egakat SPA | Control Management & Continuous Improvement | {datetime.date.today().strftime('%B %Y')}")
r.font.size=Pt(9); r.italic=True; r.font.color.rgb=RGBColor(0x80,0x80,0x80)

# Guardar
salida = r"C:\ClaudeWork\NPS_Encuesta\Propuesta_Configuracion_Encuestas_NPS_CSAT.docx"
doc.save(salida)
print(f"Documento generado: {salida}")
