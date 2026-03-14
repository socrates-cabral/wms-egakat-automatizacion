import sys
sys.stdout.reconfigure(encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Helpers ────────────────────────────────────────────────────────────────
def h1(text):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

def h2(text):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

def h3(text):
    doc.add_heading(text, level=3)

def txt(text, bold=False, italic=False, color=None):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return para

def br():
    doc.add_paragraph()

def tabla(headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:fill'), '1F6FEB')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:val'), 'clear')
        tcPr.append(shd)
    for ri, row in enumerate(rows):
        tr = t.rows[ri + 1]
        for ci, val in enumerate(row):
            tr.cells[ci].text = str(val)
        if ri % 2 == 0:
            for ci in range(len(headers)):
                tc = tr.cells[ci]._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'EBF5FB')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')
                tcPr.append(shd)
    br()

def bullet(items):
    for item in items:
        if isinstance(item, tuple):
            para = doc.add_paragraph(style='List Bullet')
            run = para.add_run(item[0] + ": ")
            run.bold = True
            para.add_run(item[1])
        else:
            doc.add_paragraph(item, style='List Bullet')

def numbered(items):
    for item in items:
        doc.add_paragraph(item, style='List Number')

# ══════════════════════════════════════════════════════════════════════════════
# PORTADA
# ══════════════════════════════════════════════════════════════════════════════
br()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Chiquito Finanzas")
run.bold = True; run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = sub.add_run("Guía completa de uso de la aplicación")
run2.font.size = Pt(16)
run2.font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

br()
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run("Para usuarios externos — Versión 1.0  |  Marzo 2026").italic = True

br()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run("Desarrollado por Sócrates Cabral\nControl de Gestión y Mejora Continua — Egakat SPA").italic = True
info.runs[0].font.color.rgb = RGBColor(0x5D, 0x6D, 0x7E)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# ÍNDICE
# ══════════════════════════════════════════════════════════════════════════════
h1("Tabla de contenidos")
secciones = [
    "1. ¿Qué es Chiquito Finanzas?",
    "2. Cómo iniciar la aplicación",
    "3. Cómo cargar los datos del negocio",
    "4. Navegación general",
    "5. Página: Dashboard — estado del negocio",
    "6. Página: Simulador de escenarios",
    "7. Página: Libro de Caja",
    "8. Página: Deuda",
    "9. Página: Plan de Acción",
    "10. Página: Inyección de Capital",
    "11. Página: Ajustes y exportar PDF",
    "12. Conceptos financieros clave",
    "13. Glosario",
    "14. Preguntas frecuentes",
]
for s in secciones:
    doc.add_paragraph(s, style='List Bullet')
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. QUÉ ES
# ══════════════════════════════════════════════════════════════════════════════
h1("1. ¿Qué es Chiquito Finanzas?")
txt("Chiquito Finanzas es una aplicación de gestión financiera diseñada específicamente para el negocio de muebles 'Chiquito', ubicado en Macul, Santiago. Permite a cualquier persona, sin conocimientos financieros avanzados, entender la salud económica del negocio, simular escenarios y tomar decisiones informadas.")
br()
txt("La aplicación lee automáticamente el archivo Excel con el libro de caja y la deuda del negocio, y convierte esos números en gráficos, alertas y recomendaciones claras.")
br()

h2("¿Para qué sirve?")
bullet([
    "Ver en tiempo real si el negocio está ganando o perdiendo dinero",
    "Entender cuánto debe el negocio y a quién",
    "Simular qué pasaría si suben las ventas, bajan los gastos o se renegocia el alquiler",
    "Decidir si conviene pedir un crédito para pagar deudas más caras",
    "Visualizar el plan de acción priorizado para salir de la crisis",
    "Generar un reporte PDF mensual para compartir con la familia o un asesor",
])
br()

h2("¿Quién la creó y sobre qué datos opera?")
txt("Fue desarrollada por Sócrates Cabral (hermano del dueño del negocio), profesional de Control de Gestión. Lee el archivo Excel 'Diagnostico_Financiero_Chiquito.xlsx' que el propio negocio mantiene actualizado con sus ingresos, gastos y deudas.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. CÓMO INICIAR
# ══════════════════════════════════════════════════════════════════════════════
h1("2. Cómo iniciar la aplicación")
txt("La aplicación se ejecuta en tu computador y se abre en el navegador web (Chrome, Edge o Firefox). No requiere conexión a internet una vez instalada.")
br()

h2("Requisitos previos")
tabla(
    ["Requisito", "Detalle"],
    [
        ["Python 3.10 o superior", "Descargar desde python.org si no está instalado"],
        ["Librerías instaladas", "Ejecutar: py -m pip install -r requirements.txt"],
        ["Archivo Excel", "Diagnostico_Financiero_Chiquito.xlsx (proporcionado)"],
        ["Navegador web", "Chrome, Edge o Firefox actualizado"],
    ]
)

h2("Pasos para iniciar")
numbered([
    "Abrir la terminal (Símbolo del sistema o PowerShell en Windows)",
    r"Ir a la carpeta del proyecto: cd C:\ClaudeWork\chiquito_financiero\app",
    "Ejecutar: py -m streamlit run main.py",
    "El navegador se abre automáticamente en http://localhost:8501",
    "La app carga el Excel automáticamente si está en la ruta configurada",
])
br()
txt("Consejo: Si no tienes el Excel en la carpeta, puedes subirlo directamente desde la app usando el botón 'Sube el Excel del mes' en el panel lateral izquierdo.", italic=True, color=(30, 130, 76))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. CARGAR DATOS
# ══════════════════════════════════════════════════════════════════════════════
h1("3. Cómo cargar los datos del negocio")
txt("La aplicación necesita el archivo Excel con el libro de caja del negocio. Existen dos formas:")
br()

h2("Opción A — Archivo local (recomendada para uso diario)")
bullet([
    "Ir a la página ⚙️ Ajustes en el menú lateral",
    r"En 'Ruta completa al Excel', escribir la ruta donde está el archivo",
    r"Ejemplo: C:\Documentos\Chiquito_Finanzas.xlsx",
    "Hacer clic en '💾 Guardar configuración'",
    "La app recordará esta ruta en el futuro",
])
br()

h2("Opción B — Subir archivo (para acceso desde otro dispositivo)")
bullet([
    "En el panel lateral, hacer clic en 'Sube el Excel del mes'",
    "Seleccionar el archivo .xlsx desde el dispositivo",
    "La app carga los datos al instante",
    "Nota: este archivo solo dura mientras la sesión esté activa",
])
br()

h2("Estructura del Excel requerida")
txt("El Excel debe tener al menos estas hojas:")
tabla(
    ["Hoja", "Contenido requerido"],
    [
        ["Cajas", "Libro de caja: columnas fecha, mes, descripción, monto (ingresos en col A-D, gastos en col J-M)"],
        ["Cajas_2026", "Mismo formato que Cajas, para el año 2026"],
        ["Deuda_2026 actual", "Tabla de deudas: acreedor, saldo, cuota, tasa"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. NAVEGACIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("4. Navegación general")
txt("La aplicación tiene un panel lateral izquierdo (sidebar) con el menú de navegación.")
br()
tabla(
    ["Página", "Ícono", "¿Para qué sirve?"],
    [
        ["Dashboard", "📊", "Resumen general: ingresos, gastos, deuda, alertas y gráficos"],
        ["Simulador", "🎛", "Probar qué pasaría si cambian ventas, alquiler u otros costos"],
        ["Libro de Caja", "💰", "Detalle de cada ingreso y gasto mes a mes con filtros"],
        ["Deuda", "💳", "Cuánto se debe, a quién y cuánto se paga por mes"],
        ["Plan de Acción", "✅", "10 acciones recomendadas en 3 horizontes de tiempo"],
        ["Inyección Capital", "💉", "Simular el impacto de pedir crédito para pagar deudas caras"],
        ["Ajustes", "⚙️", "Configurar ruta del Excel, umbrales de alerta y exportar PDF"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. DASHBOARD
# ══════════════════════════════════════════════════════════════════════════════
h1("5. Página: Dashboard — entender el estado del negocio")
txt("El Dashboard es la pantalla principal. Muestra en una sola vista la salud financiera completa del negocio.")
br()

h2("5.1 La alerta de estado (primer elemento visible)")
txt("Lo primero que verás al abrir el Dashboard es una alerta de color:")
tabla(
    ["Color", "Qué significa", "Acción recomendada"],
    [
        ["🔴 Rojo", "Resultado promedio negativo: el negocio pierde dinero cada mes", "Actuar inmediatamente — ver Plan de Acción"],
        ["🟡 Ámbar", "Ventas bajo el 70% del Punto de Equilibrio", "Atención: el negocio sobrevive pero está en riesgo"],
        ["🟢 Verde", "Ventas superan el 70% del Punto de Equilibrio", "Negocio estable — mantener y mejorar"],
    ]
)

h2("5.2 Los 6 KPIs principales")
txt("Debajo de la alerta aparecen 6 tarjetas con los indicadores más importantes:")
tabla(
    ["KPI", "Qué mide", "Valor actual", "Cómo interpretar"],
    [
        ["Ingreso prom/mes", "Promedio mensual de ingresos", "$1,842,930", "Meta: superar el PE ($4,250,896)"],
        ["Gasto prom/mes", "Promedio mensual de gastos operativos", "$1,948,700", "Debe ser menor que los ingresos"],
        ["Resultado neto prom", "Ingresos menos Gastos", "-$105,770", "Verde=ganancia, Ámbar=pérdida leve, Rojo=pérdida grave"],
        ["Deuda total", "Suma de todos los saldos pendientes", "$26,451,837", "Referencia: no superar 12 meses de ingresos"],
        ["Cuotas/mes", "Total que se paga mensualmente en deudas", "$918,903", "Ideal: máximo 30% del ingreso mensual"],
        ["% PE alcanzado", "% del Punto de Equilibrio que representan las ventas", "43%", "Meta mínima: 70%. Ideal: 100% o más"],
    ]
)

h2("5.3 Los 3 gráficos del Dashboard")
br()
h3("Gráfico 1 — Ingresos vs Gastos por mes")
txt("Barras agrupadas que comparan ingresos (verde) y gastos (rojo) mes a mes. Si las barras verdes superan a las rojas, el mes fue positivo.")
br()
h3("Gráfico 2 — Composición de costos fijos")
txt("Gráfico de dona que muestra cómo se distribuyen los costos fijos. El segmento más grande indica el gasto más crítico. En Chiquito, el alquiler representa ~38% del total de costos fijos.")
br()
h3("Gráfico 3 — Resultado neto mensual")
txt("Barras que muestran ganancia (verde, arriba de la línea cero) o pérdida (rojo, abajo) de cada mes. El objetivo es que todas las barras estén en verde sobre la línea.")
br()
h2("5.4 Indicador de actualización de datos")
txt("Bajo las alertas aparece la fecha y hora en que fue modificado por última vez el Excel. Si este dato es muy antiguo, los KPIs podrían no reflejar la situación actual.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 6. SIMULADOR
# ══════════════════════════════════════════════════════════════════════════════
h1("6. Página: Simulador de escenarios")
txt("El Simulador permite responder preguntas del tipo '¿qué pasaría si...?' sin afectar los datos reales.")
br()

h2("6.1 Escenarios rápidos")
tabla(
    ["Botón", "Qué simula", "Ventas", "Alquiler"],
    [
        ["📍 Actual", "Estado real del negocio hoy", "$1,860,000", "$700,000"],
        ["🌱 Optimista", "Si las ventas suben a $3,000,000", "$3,000,000", "$700,000"],
        ["🤝 Renegociado", "Si se negocia alquiler y TCs", "$2,200,000", "$450,000"],
        ["⚖️ Equilibrio", "Ventas exactas para cubrir todo", "$3,950,000", "$700,000"],
    ]
)

h2("6.2 Parámetros ajustables con sliders")
bullet([
    ("Ventas objetivo/mes", "Cuánto dinero entraría por ventas cada mes"),
    ("Alquiler taller", "Cuánto se paga de arriendo mensual al taller"),
    ("Cuotas TC", "Cuánto se paga en total por las tarjetas de crédito"),
    ("Margen bruto (%)", "Qué porcentaje de cada venta queda como ganancia bruta"),
    ("Crecimiento mensual", "Si las ventas crecen un porcentaje fijo cada mes"),
])
br()

h2("6.3 Cómo leer los resultados")
tabla(
    ["KPI del Simulador", "Qué muestra"],
    [
        ["Resultado mensual", "Si el negocio gana o pierde con ese escenario (verde=ganancia)"],
        ["Punto de equilibrio", "Las ventas mínimas para ese escenario específico"],
        ["% PE alcanzado", "Qué tan cerca están las ventas de la meta"],
    ]
)
br()
txt("Consejo de uso: Ajusta los sliders hasta que el resultado mensual sea positivo y sostenible. Eso te da la 'meta mínima' que el negocio debe alcanzar para sobrevivir.", italic=True, color=(30, 130, 76))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 7. LIBRO DE CAJA
# ══════════════════════════════════════════════════════════════════════════════
h1("7. Página: Libro de Caja")
txt("Muestra el detalle completo de cada movimiento del negocio: cada ingreso por venta y cada gasto pagado, leídos directamente desde el Excel.")
br()
txt("Usa los filtros para responder preguntas como: ¿cuánto gané en enero? ¿cuánto gasté en arriendo este año? ¿qué meses fueron rentables?")
br()
tabla(
    ["Filtro", "Cómo usarlo"],
    [
        ["Filtrar por mes", "Selecciona uno o varios meses del menú desplegable"],
        ["Tipo", "Elige 'ingreso', 'gasto' o ambos"],
    ]
)
br()
txt("Las 3 métricas sobre la tabla (ingresos, gastos y resultado) se recalculan en tiempo real según los filtros aplicados.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 8. DEUDA
# ══════════════════════════════════════════════════════════════════════════════
h1("8. Página: Deuda")
txt("Muestra el estado completo de todas las deudas del negocio. Es una de las páginas más importantes dado el nivel de endeudamiento actual.")
br()

h2("8.1 Métricas superiores")
tabla(
    ["Métrica", "Valor actual", "Qué significa"],
    [
        ["Deuda total consolidada", "$26,451,837", "Suma de todos los saldos pendientes con bancos, retail y familia"],
        ["Cuotas/mes consolidadas", "$918,903", "Lo que se paga cada mes entre todas las deudas — equivale al 50% del ingreso promedio"],
        ["Instrumentos activos", "8", "Cantidad de deudas distintas activas en este momento"],
    ]
)

h2("8.2 Desglose por categoría de deuda")
tabla(
    ["Categoría", "Acreedores", "Tasa aprox.", "Riesgo"],
    [
        ["Créditos bancarios", "Itaú ($5.7M) + Banco Estado ($5.6M)", "2.8–3.1%/mes", "Medio"],
        ["Tarjetas de crédito", "Santander ($3.7M) + CMR Falabella ($1.6M)", "2.8–3.3%/mes", "Alto — pueden entrar en mora"],
        ["Líneas de crédito", "3 bancos ($2.3M)", "3.1%/mes", "Alto — muy costosas"],
        ["Automotriz", "Crédito camión Foton ($9.5M)", "1.2%/mes", "Bajo — respaldado por el activo"],
        ["Familiar", "Hermana ($1.0M en dólares)", "0%/mes", "Ninguno — sin interés ni cuota"],
    ]
)

h2("8.3 Gráfico de barras de deuda")
txt("Barras horizontales ordenadas de mayor a menor saldo. El camión Foton es el mayor saldo ($9.5M) pero tiene la tasa más baja. Las tarjetas, aunque de menor saldo, son las más dañinas por su alta tasa de interés y riesgo de mora.")
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 9. PLAN DE ACCIÓN
# ══════════════════════════════════════════════════════════════════════════════
h1("9. Página: Plan de Acción")
txt("Presenta 10 acciones concretas en 3 horizontes de tiempo, ordenadas por urgencia e impacto en el flujo mensual del negocio.")
br()

h2("Horizonte 1 — URGENTE (0 a 3 meses)")
tabla(
    ["#", "Acción", "Ahorro estimado/mes"],
    [
        ["1", "Renegociar alquiler del taller de $700,000 a $400,000", "$300,000"],
        ["2", "Separar gastos personales de la caja del negocio", "$80,000"],
        ["3", "Renegociar tarjetas de crédito antes de entrar en mora formal", "$150,000"],
        ["4", "Subir precios un 10-15% en todos los productos", "$186,000"],
    ]
)

h2("Horizonte 2 — IMPORTANTE (3 a 12 meses)")
tabla(
    ["#", "Acción", "Impacto esperado"],
    [
        ["5", "Escalar ventas a $3,500,000/mes", "Meta mínima de supervivencia del negocio"],
        ["6", "Formalizar la empresa como EIRL o SpA", "Acceso a créditos SERCOTEC y FOGAPE (tasas 0.5-1%/mes)"],
        ["7", "Evaluar vender el camión Foton si ventas no suben", "Libera $329,778/mes en cuota + seguro"],
    ]
)

h2("Horizonte 3 — LARGO PLAZO (12 a 36 meses)")
tabla(
    ["#", "Acción", "Impacto esperado"],
    [
        ["8", "Liquidar tarjetas de crédito (CMR primero)", "Libera flujo permanente y mejora historial crediticio"],
        ["9", "Construir colchón de liquidez de 2 meses (~$3,800,000)", "Estabilidad ante imprevistos y temporadas bajas"],
        ["10", "Contratar vendedor/a a comisión (sin sueldo fijo)", "Multiplica ventas sin aumentar costos fijos"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 10. INYECCIÓN CAPITAL
# ══════════════════════════════════════════════════════════════════════════════
h1("10. Página: Inyección de Capital")
txt("Simula el impacto de pedir un crédito bancario para pagar las deudas más caras del negocio, aprovechando el arbitraje de tasa (reemplazar deuda al 3.1-3.3%/mes por crédito al 1.43%/mes).")
br()

h2("10.1 Comparador de opciones BCI")
txt("Muestra 3 opciones reales de crédito BCI por $10,000,000:")
tabla(
    ["Opción", "Cuotas", "Tasa/mes", "Cuota mensual", "Costo total", "Recomendada"],
    [
        ["A", "18", "1.43%", "$648,805", "$11,678,482", "No"],
        ["B", "24", "1.51%", "$505,611", "$12,134,664", "Sí — menor cuota mensual"],
        ["C", "24", "1.40%", "$508,179", "$12,196,296", "No"],
    ]
)
txt("La Opción B se recomienda porque aunque paga $456,182 más en total, la cuota mensual es $143,194 más baja. Para un negocio con flujo de caja ajustado, la liquidez mensual es prioritaria.", italic=True, color=(30, 130, 76))
br()

h2("10.2 Cómo se asigna el capital")
txt("La app asigna el capital disponible pagando primero las deudas con mayor tasa de interés:")
numbered([
    "CMR Falabella — $1,607,443 al 3.3%/mes (la más cara) → libera $80,000/mes",
    "Líneas de crédito — $2,360,000 al 3.1%/mes → libera $71,660/mes",
    "Banco Estado — $5,600,000 al 3.1%/mes → libera $174,437/mes",
    "Santander — pago parcial con lo que reste del capital disponible",
])
br()

h2("10.3 KPIs clave de la simulación")
tabla(
    ["KPI", "Valor típico", "Qué significa"],
    [
        ["Cuotas liberadas", "$326,097/mes", "Lo que se deja de pagar al eliminar 3 deudas completas"],
        ["Nueva cuota BCI", "$505,611/mes", "Lo nuevo que hay que pagar por el crédito BCI"],
        ["Impacto neto corto plazo", "-$179,514/mes", "La diferencia: se paga más al inicio pero con mejor estructura"],
        ["Ahorro en intereses/mes", "+$280,268/mes", "Lo que se ahorra cada mes al bajar la tasa promedio"],
        ["Ahorro total (24 meses)", "+$3,363,000", "Ahorro acumulado durante todo el período del crédito"],
        ["Arbitraje de tasa", "1.7%/mes", "Ventaja por reemplazar deuda cara por deuda barata"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 11. AJUSTES Y PDF
# ══════════════════════════════════════════════════════════════════════════════
h1("11. Página: Ajustes y exportar PDF")
br()
h2("Configuración de la ruta del Excel")
txt("Escribe la ruta completa al archivo Excel y haz clic en 'Guardar configuración'. La app recordará esta ruta para futuros usos.")
br()

h2("Generar reporte PDF mensual")
numbered([
    "Seleccionar el mes del reporte en el selector",
    "Hacer clic en '📄 Generar PDF'",
    "Hacer clic en '⬇️ Descargar PDF' cuando aparezca el botón",
    "El PDF incluye KPIs del mes, comparación con el PE, estado de deuda y resultado neto",
])
br()
txt("Recomendación: generar y guardar el PDF al cierre de cada mes como respaldo histórico.", italic=True, color=(30, 130, 76))
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 12. CONCEPTOS FINANCIEROS
# ══════════════════════════════════════════════════════════════════════════════
h1("12. Conceptos financieros clave")
br()

conceptos = [
    ("Punto de Equilibrio (PE)",
     "El nivel mínimo de ventas que necesita el negocio para no perder dinero. Si las ventas igualan el PE, el resultado es exactamente cero. Por encima, hay ganancia; por debajo, hay pérdida.",
     "En Chiquito: PE actual = $4,250,896/mes. Ventas actuales = $1,842,930/mes (43% del PE). El negocio necesita más que duplicar sus ventas para cubrir todos sus costos."),
    ("Margen bruto",
     "El porcentaje de cada peso vendido que queda después de pagar los materiales. Un margen del 45% significa que de cada $100 vendidos, $45 quedan para pagar costos fijos y deudas.",
     "En Chiquito: margen promedio estimado 42-50% según el producto. La estantería estándar blanca tiene el mejor margen: 55.2%."),
    ("Costos fijos",
     "Gastos que se pagan siempre, independientemente de cuánto se venda. No desaparecen aunque el negocio no venda nada en un mes.",
     "En Chiquito: alquiler $700,000 + servicios $118,000 + gasolina $100,000 + otros = $994,000/mes solo en costos fijos operativos, más $918,903 en cuotas bancarias."),
    ("Tasa de interés mensual",
     "El porcentaje que cobra el banco o la entidad financiera por el dinero prestado, cobrado cada mes sobre el saldo pendiente.",
     "Ejemplo: tarjeta con saldo $1,000,000 al 3%/mes genera $30,000 en intereses cada mes, además del pago al capital. Cuanto más alta la tasa, más caro es tener esa deuda."),
    ("Arbitraje de tasa",
     "Estrategia que consiste en tomar deuda a tasa baja para pagar deuda a tasa alta, reduciendo el costo financiero total sin cambiar el monto de la deuda.",
     "En Chiquito: reemplazar deudas al 3.1-3.3%/mes por crédito BCI al 1.43%/mes genera un ahorro neto de ~$280,000/mes en intereses."),
    ("Flujo de caja",
     "El movimiento real de dinero que entra y sale del negocio. Diferente de las ganancias contables: el flujo de caja mide cuándo y cuánto dinero efectivo está disponible para operar.",
     "El Libro de Caja en la app registra exactamente esto: cada cobro a un cliente y cada pago a un proveedor o banco."),
]

for nombre, definicion, ejemplo in conceptos:
    h2(nombre)
    txt(definicion)
    br()
    para = doc.add_paragraph()
    run = para.add_run("Ejemplo práctico: ")
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    para.add_run(ejemplo)
    br()

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 13. GLOSARIO
# ══════════════════════════════════════════════════════════════════════════════
h1("13. Glosario")
tabla(
    ["Término", "Definición"],
    [
        ["CAE", "Carga Anual Equivalente — costo total del crédito expresado como % anual, incluye todos los costos y comisiones"],
        ["CMR", "Tarjeta de crédito retail de Falabella — tasa ~3.3%/mes"],
        ["CTC", "Costo Total del Crédito — suma total de todas las cuotas a pagar en la vida del crédito"],
        ["EIRL", "Empresa Individual de Responsabilidad Limitada — separa patrimonio personal del negocio"],
        ["FOGAPE", "Fondo de Garantía para Pequeños Empresarios — aval estatal para créditos PYME"],
        ["KPI", "Key Performance Indicator — indicador clave para medir el desempeño del negocio"],
        ["PE", "Punto de Equilibrio — ventas mínimas para no tener ni ganancia ni pérdida"],
        ["SERCOTEC", "Servicio de Cooperación Técnica — entrega microcréditos y subsidios a pequeños negocios en Chile"],
        ["SpA", "Sociedad por Acciones — figura legal chilena más flexible que la EIRL"],
        ["TC", "Tarjeta de Crédito"],
        ["TMC", "Tasa Máxima Convencional — tasa máxima legal en Chile (~2.75%/mes en 2026)"],
    ]
)
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 14. FAQ
# ══════════════════════════════════════════════════════════════════════════════
h1("14. Preguntas frecuentes")
br()
faqs = [
    ("¿Los datos son privados?",
     "Sí. La app corre completamente en tu computador. Ningún dato se envía a internet."),
    ("¿Puedo usarla desde el celular?",
     "Sí, si la app corre en un computador de la misma red WiFi, accede desde el celular usando la 'Network URL' que aparece al iniciarla (ej: http://192.168.x.x:8501)."),
    ("¿Qué pasa si cierro la app?",
     "Los datos del Excel no se pierden. Al reiniciar, todo vuelve a cargar desde el archivo."),
    ("¿Con qué frecuencia actualizar el Excel?",
     "Al menos una vez por semana registrando todos los ingresos y gastos del período."),
    ("¿Qué significa que las ventas estén al 43% del PE?",
     "Que el negocio vende menos de la mitad de lo que necesita para cubrir todos sus costos. Si esta situación no mejora, el negocio seguirá perdiendo dinero cada mes."),
    ("¿El PDF se guarda automáticamente?",
     "No. Debes ir a Ajustes, seleccionar el mes y hacer clic en 'Generar PDF', luego descargarlo. Hazlo al cierre de cada mes."),
    ("¿Puedo modificar los costos fijos?",
     "Los costos fijos base están en el código (calculators.py). Para modificarlos definitivamente hay que editar ese archivo. En el Simulador puedes ajustar el alquiler de forma temporal."),
]

for pregunta, respuesta in faqs:
    para = doc.add_paragraph()
    run = para.add_run("P: " + pregunta)
    run.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x6F, 0xEB)
    para2 = doc.add_paragraph()
    para2.add_run("R: " + respuesta)
    br()

doc.add_page_break()
txt("Chiquito Finanzas v1.0  —  Desarrollado por Sócrates Cabral, Control de Gestión y Mejora Continua, Egakat SPA  —  Marzo 2026", italic=True, color=(93, 109, 126))

output = r"C:\ClaudeWork\chiquito_financiero\Guia_ChiquitoFinanzas.docx"
doc.save(output)
print(f"[OK] Word guardado: {output}")
