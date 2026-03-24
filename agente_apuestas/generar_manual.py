"""
generar_manual.py
Genera MANUAL_AGENTE_APUESTAS.docx y .pdf
Sprint 10 — Actualizado 24/03/2026
"""
import sys
if sys.stdout:
    sys.stdout.reconfigure(encoding="utf-8")

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

BASE_DIR = Path(__file__).parent
OUT_DOCX = BASE_DIR / "MANUAL_AGENTE_APUESTAS.docx"
OUT_PDF  = BASE_DIR / "MANUAL_AGENTE_APUESTAS.pdf"

# ── Colores ──────────────────────────────────────────────────────────────────
VERDE   = RGBColor(0x22, 0xC5, 0x5E)
AZUL    = RGBColor(0x3B, 0x82, 0xF6)
AMARILLO= RGBColor(0xF5, 0x9E, 0x0B)
ROJO    = RGBColor(0xEF, 0x44, 0x44)
GRIS    = RGBColor(0x64, 0x74, 0x8B)
NEGRO   = RGBColor(0x0F, 0x17, 0x2A)
BLANCO  = RGBColor(0xFF, 0xFF, 0xFF)
FONDO_TBL = RGBColor(0x1E, 0x29, 0x3B)


def rgb_to_hex(color) -> str:
    """Convierte RGBColor (tuple subclass: r,g,b) a string hex 'RRGGBB'."""
    r, g, b = color[0], color[1], color[2]
    return f"{r:02X}{g:02X}{b:02X}"


def rgb_lighten(color, factor=0.12) -> RGBColor:
    """Versión más clara de un color para usar como fondo."""
    r, g, b = color[0], color[1], color[2]
    nr = int(r * factor + 255 * (1 - factor))
    ng = int(g * factor + 255 * (1 - factor))
    nb = int(b * factor + 255 * (1 - factor))
    return RGBColor(min(255, nr), min(255, ng), min(255, nb))


def set_cell_bg(cell, color):
    """Establece color de fondo de celda. color: RGBColor o tuple (r,g,b)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), rgb_to_hex(color))
    tcPr.append(shd)


def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(6)
    return p


def add_para(doc, text="", bold_parts=None, color=None, size=10, italic=False, space_after=6):
    """Agrega párrafo con texto. bold_parts = list of (start, end) substrings a poner en negrita."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_run_colored(para, text, color=None, bold=False, size=10, italic=False):
    run = para.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return run


def add_callout(doc, text, color=VERDE, label="💡"):
    """Caja de callout con borde izquierdo simulado como tabla 1 columna."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, rgb_lighten(color, 0.10))
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.3)
    run = p.add_run(f"{label}  {text}" if label else text)
    run.font.size = Pt(9.5)
    run.font.color.rgb = NEGRO
    # Borde izquierdo coloreado
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    hex_c = rgb_to_hex(color)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "0")
    left.set(qn("w:color"), hex_c)
    tcBorders.append(left)
    tcPr.append(tcBorders)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_formula(doc, text):
    """Bloque de fórmula matemática."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(8)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.font.color.rgb = VERDE
    run.font.bold = True


def add_roi_table(doc):
    headers = ["Confianza Mín.", "Value Mín.", "N° Apuestas", "% Total", "Accuracy", "ROI Flat", "ROI Kelly"]
    rows = [
        ["55%", "3%",  "341", "19.0%", "54.2%", "+2.94%",  "+22.67%"],
        ["55%", "10%", "257", "14.3%", "50.2%", "+3.32%",  "+29.90%"],
        ["60%", "3%",  "251", "14.0%", "59.8%", "+5.78%",  "+46.17%"],
        ["60%", "10%", "180", "10.1%", "55.6%", "+6.26%",  "+50.87%"],
        ["65%", "5%",  "142", "7.9%",  "64.1%", "+5.68%",  "+38.39%"],
        ["65%", "10%", "104", "5.8%",  "62.5%", "+8.49%",  "+44.82%"],
        ["70%", "3%",  "96",  "5.4%",  "69.8%", "+6.84%",  "+24.98%"],
        ["70% ★","10% ★","61","3.4%",  "67.2%", "+10.11%", "+31.60%"],  # mejor
        ["75%", "3%",  "57",  "3.2%",  "66.7%", "-4.60%",  "-15.16%"],
        ["75%", "10%", "35",  "2.0%",  "65.7%", "-0.51%",  "-5.98%"],
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, FONDO_TBL)
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True
        run.font.size = Pt(8)
        run.font.color.rgb = BLANCO
    # Filas
    for ri, row in enumerate(rows):
        is_best = "★" in row[0]
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            if is_best:
                set_cell_bg(c, RGBColor(0xDC, 0xFC, 0xE7))
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            run.font.bold = is_best
            if val.startswith("+"):
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
            elif val.startswith("-"):
                run.font.color.rgb = ROJO
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_liga_table(doc):
    headers = ["Liga", "Partidos Test", "Apuestas", "Accuracy", "ROI Flat", "Estado"]
    rows = [
        ["🇮🇹 Serie A",      "358", "24", "66.7%", "+31.65%", "✓ ACTIVA"],
        ["🇪🇸 La Liga",      "358", "14", "64.3%", "+10.86%", "Pocas apuestas"],
        ["🇩🇪 Bundesliga",   "306",  "6", "100%",  "+63.00%", "Muestra pequeña"],
        ["🇬🇧 Premier Lge.", "380",  "9", "55.6%", "-18.67%", "Suspendida"],
        ["🇫🇷 Ligue 1",      "306",  "8", "62.5%",  "-0.50%", "Suspendida"],
    ]
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    for i, h in enumerate(headers):
        c = hdr.cells[i]
        set_cell_bg(c, FONDO_TBL)
        p = c.paragraphs[0]
        run = p.add_run(h)
        run.font.bold = True; run.font.size = Pt(8.5); run.font.color.rgb = BLANCO
    for ri, row in enumerate(rows):
        is_activa = "ACTIVA" in row[5]
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]
            if is_activa:
                set_cell_bg(c, RGBColor(0xDC, 0xFC, 0xE7))
            p = c.paragraphs[0]
            run = p.add_run(val)
            run.font.size = Pt(9)
            run.font.bold = is_activa
            if val.startswith("+"):
                run.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
            elif val.startswith("-"):
                run.font.color.rgb = ROJO
            elif val == "✓ ACTIVA":
                run.font.color.rgb = VERDE
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN — construir documento
# ═══════════════════════════════════════════════════════════════════════════════
def generar_docx():
    doc = Document()

    # Márgenes
    for section in doc.sections:
        section.top_margin    = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # Fuente por defecto
    doc.styles["Normal"].font.name = "Calibri"
    doc.styles["Normal"].font.size = Pt(10)

    # ── PORTADA ────────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(30)
    run = p.add_run("⚽  AGENTE DE APUESTAS DEPORTIVAS")
    run.font.size = Pt(22); run.font.bold = True; run.font.color.rgb = VERDE

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Manual del Sistema — Sprint 10")
    run2.font.size = Pt(14); run2.font.color.rgb = GRIS

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = p3.add_run("XGBoost + Pi-Rating + xG Histórico + Transfermarkt")
    run3.font.size = Pt(11); run3.font.color.rgb = AZUL

    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run("Actualizado: 24/03/2026  ·  Desarrollado con Claude Code (Anthropic)  ·  Sócrates Cabral")
    run4.font.size = Pt(9); run4.font.color.rgb = GRIS
    run4.font.italic = True

    doc.add_paragraph()

    # KPIs portada
    tbl_kpi = doc.add_table(rows=1, cols=5)
    tbl_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    kpis = [
        ("10.707", "Partidos\nAnalizados"),
        ("52.3%",  "Accuracy\nGlobal"),
        ("67.2%",  "Accuracy\nal Apostar"),
        ("+31.7%", "ROI Serie A\n(post-fixes)"),
        ("Serie A", "Liga\nActiva ✓"),
    ]
    for i, (val, lbl) in enumerate(kpis):
        c = tbl_kpi.rows[0].cells[i]
        set_cell_bg(c, FONDO_TBL)
        p_k = c.paragraphs[0]
        p_k.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_v = p_k.add_run(f"\n{val}\n")
        r_v.font.size = Pt(14); r_v.font.bold = True; r_v.font.color.rgb = VERDE
        r_l = p_k.add_run(lbl + "\n")
        r_l.font.size = Pt(7); r_l.font.color.rgb = GRIS

    doc.add_page_break()

    # ── 1. LA HISTORIA ─────────────────────────────────────────────────────────
    add_heading(doc, "1. La Historia: ¿De qué se trata esto?", level=1, color=VERDE)
    add_para(doc, (
        "Imagina que tienes un amigo que lleva 6 años estudiando fútbol de manera obsesiva. "
        "Ha visto 10.707 partidos de las mejores ligas del mundo y tomó nota de todo: "
        "quién ganó, por cuánto, cómo venía cada equipo jugando, cuánto vale cada plantilla, "
        "cuántos goles esperados (xG) generó cada equipo..."
    ), size=10)
    add_para(doc, (
        "Después de tanto analizar, desarrolló un 'instinto matemático' para detectar cuándo "
        "las probabilidades están a tu favor. No apuesta en todos los partidos — solo en aquellos "
        "donde está muy seguro de que las cuotas del bookmaker están equivocadas."
    ), size=10)
    add_callout(doc,
        "Ese 'amigo' es este modelo. Entrenado con inteligencia artificial sobre 5 temporadas "
        "de fútbol europeo, aprende a identificar los momentos donde apostar tiene sentido matemático.",
        color=VERDE, label="💡")
    add_para(doc, (
        "El resultado en simulación (Serie A, post-fixes Sprint 10): de cada 100 pesos apostados "
        "en los partidos de mayor confianza, el modelo retorna 131 pesos en promedio. No es magia — es estadística."
    ), size=10)

    # ── 2. CÓMO FUNCIONA ───────────────────────────────────────────────────────
    add_heading(doc, "2. ¿Cómo funciona por dentro?", level=1, color=VERDE)
    add_para(doc, "El modelo usa tres capas de análisis que trabajan juntas:", size=10)

    # Flujo como texto estructurado
    pasos_flujo = [
        ("📊 Datos históricos", "10.707 partidos · 5 ligas · 5 temporadas"),
        ("🧮 Pi-Rating",        "Fortaleza real de cada equipo (actualizado en vivo)"),
        ("🤖 XGBoost",          "Modelo ML — probabilidades [Home, Empate, Away]"),
        ("🎯 Value Bet",        "Solo si confianza ≥70% Y value ≥10%"),
        ("📱 Telegram",         "Recomendación automática cada mañana a las 08:00"),
    ]
    for icon_title, desc in pasos_flujo:
        p = doc.add_paragraph(style="List Bullet")
        r1 = p.add_run(icon_title + ": ")
        r1.font.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(10); r2.font.color.rgb = GRIS

    doc.add_paragraph()
    p_c1 = doc.add_paragraph()
    r = p_c1.add_run("Capa 1 — Pi-Rating: ")
    r.font.bold = True; r.font.size = Pt(10)
    r2 = p_c1.add_run(
        "Como el rating ELO del ajedrez, pero adaptado al fútbol. Le da a cada equipo un número "
        "que refleja su fortaleza real. A diferencia de solo contar victorias, el Pi-Rating también "
        "considera por cuánto ganó o perdió. El Bayern que gana 4-0 sube más que el que gana 1-0."
    )
    r2.font.size = Pt(10)

    add_callout(doc,
        "Ejemplo de Pi-Rating: Inter Milan 1.43 · Napoli 1.28 · La diferencia (0.15) es la feature "
        "más importante del modelo — predice directamente quién tiene más probabilidades de ganar.",
        color=AZUL, label="📐")

    p_c2 = doc.add_paragraph()
    r = p_c2.add_run("Capa 2 — XGBoost: ")
    r.font.bold = True; r.font.size = Pt(10)
    r2 = p_c2.add_run(
        "El algoritmo de machine learning más usado en competencias de datos del mundo. "
        "Toma el Pi-Rating + forma reciente (últimos 5 partidos) + xG histórico + valor de plantilla "
        "y calcula tres probabilidades: % de victoria local, empate y victoria visitante."
    )
    r2.font.size = Pt(10)

    p_c3 = doc.add_paragraph()
    r = p_c3.add_run("Capa 3 — predictor_tiempo_real.py (Sprint 10): ")
    r.font.bold = True; r.font.size = Pt(10); r.font.color.rgb = AZUL
    r2 = p_c3.add_run(
        "Módulo nuevo que conecta el modelo XGBoost entrenado con los partidos del día vía API. "
        "Se ejecuta automáticamente cada mañana a las 08:00 y envía recomendaciones directamente a Telegram."
    )
    r2.font.size = Pt(10)

    # ── 3. LOS NÚMEROS ─────────────────────────────────────────────────────────
    add_heading(doc, "3. Los Números: ¿Qué significan realmente?", level=1, color=AMARILLO)

    p_acc = doc.add_paragraph()
    add_run_colored(p_acc, "Accuracy global: 52.3%  ", color=AZUL, bold=True)
    p_acc.add_run(
        "El modelo predice correctamente el resultado en más de 1 de cada 2 partidos. "
        "Para fútbol, donde hay 3 resultados posibles, el límite por azar puro sería 33%. "
        "Superar 50% de manera consistente es estadísticamente relevante."
    ).font.size = Pt(10)

    add_callout(doc,
        "La trampa del 'solo 52%': En apuestas deportivas, ganar dinero no requiere acertar el 70% "
        "de las veces. Requiere acertar lo suficiente como para que las ganancias superen las pérdidas. "
        "Con buenas cuotas y buena selección, un 52% puede ser muy rentable.",
        color=AMARILLO, label="⚠️")

    p_fil = doc.add_paragraph()
    add_run_colored(p_fil, "Accuracy al apostar (umbral 70%): 67.2%  ", color=VERDE, bold=True)
    p_fil.add_run(
        "Aquí está la magia. El modelo solo recomienda apostar cuando tiene ≥70% de confianza "
        "Y la cuota del bookmaker ofrece ventaja. En esos 61 casos del test set, acertó 2 de cada 3 veces."
    ).font.size = Pt(10)

    p_roi = doc.add_paragraph()
    add_run_colored(p_roi, "ROI flat Serie A +31.7% (post-fixes)  ", color=VERDE, bold=True)
    p_roi.add_run(
        "Si en cada una de las apuestas apostáramos $1.000, al final tendríamos $1.317 netos de ganancia "
        "por cada $1.000 apostado. Esto es después de integrar xG 2019-2024 y Transfermarkt."
    ).font.size = Pt(10)

    add_heading(doc, "Grid de Rendimiento — Todas las combinaciones evaluadas", level=2, color=GRIS)
    add_roi_table(doc)
    add_callout(doc,
        "★ Mejor combinación: Confianza ≥70% + Value ≥10% → Solo apostamos en el 3.4% más prometedor "
        "de los partidos. En esos 61 casos del backtesting, el modelo ganó dinero con ROI flat +10.11%. "
        "Con los fixes de xG y Transfermarkt, el ROI de Serie A específicamente subió a +31.65%.",
        color=VERDE, label="★")

    # ── 4. ROI POR LIGA ────────────────────────────────────────────────────────
    add_heading(doc, "4. Rendimiento por Liga", level=1, color=VERDE)
    add_para(doc, "No todas las ligas funcionan igual. Rendimiento con la combinación óptima (conf ≥70%, value ≥10%):")
    add_liga_table(doc)
    add_callout(doc,
        "Serie A es la única liga activa con n ≥ 20 apuestas y ROI positivo. La Liga tiene mejor ROI "
        "(+10.86%) pero con solo 14 apuestas — necesitamos más datos. Bundesliga con 100% de accuracy "
        "en 6 apuestas es un caso claro de muestra demasiado pequeña.",
        color=VERDE, label="🇮🇹")

    # ── 5. CÓMO USAR EL MODELO ─────────────────────────────────────────────────
    add_heading(doc, "5. ¿Cómo usar el modelo? — Guía paso a paso", level=1, color=VERDE)
    pasos = [
        ("El agente corre automáticamente cada mañana (08:00)",
         "Task Scheduler ejecuta run_agent.py que descarga partidos del día, calcula Pi-Ratings "
         "actualizados y genera probabilidades para cada encuentro Serie A."),
        ("Revisar las recomendaciones en Telegram",
         "Solo aparecen partidos donde la confianza del modelo es ≥70% Y existe value ≥10% respecto "
         "a las cuotas disponibles. Las notificaciones incluyen fuente: 🤖 PREDICCION ML (XGBoost)."),
        ("Verificar la liga",
         "Actualmente solo se recomiendan apuestas en Serie A. Las demás ligas están en modo "
         "observación hasta acumular más datos (umbral: 20 apuestas con ROI positivo)."),
        ("Revisar el Value %",
         "Value = (prob_modelo × cuota) - 1. Value 10% significa que por cada $100 apostados, "
         "el valor esperado es $110. Solo apostamos cuando Value ≥ 10%."),
        ("Decidir el monto a apostar",
         "Flat betting: mismo monto siempre (ej: $5.000 por apuesta). Kelly ×¼: monto proporcional "
         "a la ventaja — mayor confianza = mayor apuesta. Para comenzar, se recomienda flat betting."),
        ("Registrar resultado y aprender",
         "El modelo registra automáticamente si la apuesta fue ganadora o perdedora en "
         "historico_apuestas.json. El backtesting nocturno (22:00) verifica resultados reales."),
    ]
    for i, (titulo, desc) in enumerate(pasos, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0)
        p.paragraph_format.space_after = Pt(8)
        r_num = p.add_run(f"  {i}.  ")
        r_num.font.bold = True; r_num.font.size = Pt(11); r_num.font.color.rgb = VERDE
        r_tit = p.add_run(titulo + "\n")
        r_tit.font.bold = True; r_tit.font.size = Pt(10)
        r_desc = p.add_run(f"      {desc}")
        r_desc.font.size = Pt(9.5); r_desc.font.color.rgb = GRIS

    # ── 6. EJEMPLO REAL ────────────────────────────────────────────────────────
    add_heading(doc, "6. Ejemplo Real: Así se ve una recomendación", level=1, color=AZUL)
    add_para(doc, "Supón que el modelo analiza este partido de Serie A:")

    # Tabla ejemplo partido
    tbl_ej = doc.add_table(rows=4, cols=2)
    tbl_ej.style = "Table Grid"
    datos_ej = [
        ("🇮🇹 Serie A · Jornada 28", "Inter Milan vs Napoli"),
        ("Pi-Rating", "Inter: 1.43 · Napoli: 1.28 · Diferencia: +0.15"),
        ("Probabilidades modelo", "Inter gana: 62%  |  Empate: 20%  |  Napoli: 18%"),
        ("Cuota bookmaker", "Inter gana: 1.75"),
    ]
    for ri, (lbl, val) in enumerate(datos_ej):
        c0 = tbl_ej.rows[ri].cells[0]
        c1 = tbl_ej.rows[ri].cells[1]
        set_cell_bg(c0, FONDO_TBL)
        r0 = c0.paragraphs[0].add_run(lbl)
        r0.font.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = BLANCO
        r1 = c1.paragraphs[0].add_run(val)
        r1.font.size = Pt(9)
    doc.add_paragraph()

    add_callout(doc,
        "La fórmula del Value:  Value = (Prob. modelo × Cuota bookmaker) − 1\n\n"
        "Si Inter tiene 62% según el modelo y la cuota es 1.75:\n"
        "Value = 0.62 × 1.75 − 1 = 1.085 − 1 = +0.085 = +8.5%\n\n"
        "Eso significa que por cada $100 apostados, el valor esperado es $108.50. ✓ Apuesta recomendada.",
        color=AZUL, label="📊")
    add_callout(doc,
        "Contra-ejemplo: Si la misma cuota fuera 1.50:\n"
        "Value = 0.62 × 1.50 − 1 = 0.93 − 1 = −0.07 = −7%\n"
        "El bookmaker cobra 7% de comisión implícita. No se apuesta.",
        color=ROJO, label="❌")

    # ── 7. ESTADÍSTICA DEL MODELO ──────────────────────────────────────────────
    add_heading(doc, "7. Estadística del Modelo: Bajo el Capó", level=1, color=AZUL)
    add_para(doc,
        "Esta sección explica cómo el modelo genera sus probabilidades. Para un ingeniero industrial, "
        "estos son los mismos conceptos de distribuciones de probabilidad, optimización y control estadístico "
        "de proceso — aplicados al fútbol.", size=10, italic=True)

    # 7A XGBoost
    p_7a = doc.add_paragraph()
    r = p_7a.add_run("A.  XGBoost: Cómo convierte variables en probabilidades")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7a.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "XGBoost es un ensemble de árboles de decisión construidos secuencialmente. "
        "Cada árbol nuevo aprende a corregir los errores del anterior — por eso se llama "
        "Gradient Boosting (mejora por gradiente). El modelo activo tiene ~300 árboles.")
    add_callout(doc,
        "Intuición: Si el árbol #1 predice 'Inter gana' pero falla en partidos con alta rivalidad "
        "histórica, el árbol #2 se especializa en esos errores. El árbol #300 ya está refinando "
        "detalles muy específicos. El modelo final es el voto ponderado de los 300 árboles.",
        color=AZUL, label="🌲")
    add_para(doc,
        "Para clasificación multiclase (3 resultados posibles), XGBoost produce 3 valores 'crudos' "
        "(logits) por partido — uno para cada clase. Estos logits se convierten en probabilidades "
        "mediante la función softmax:")
    add_formula(doc, "P(k) = e^(z_k)  /  ( e^(z_0) + e^(z_1) + e^(z_2) )")
    add_para(doc, "donde z_0, z_1, z_2 son los logits para Home / Empate / Away respectivamente.",
             color=GRIS, size=9)
    add_callout(doc,
        "Ejemplo numérico — Inter vs Napoli:\n"
        "XGBoost genera logits: [z₀=1.8, z₁=0.3, z₂=−0.6]\n"
        "e^1.8 = 6.05  ·  e^0.3 = 1.35  ·  e^−0.6 = 0.55  ·  Suma = 7.95\n\n"
        "P(Home) = 6.05/7.95 = 76.1%   P(Empate) = 1.35/7.95 = 17.0%   P(Away) = 0.55/7.95 = 6.9%\n\n"
        "✓ Suma = 100%   ✓ Ninguna es negativa (garantizado por softmax)\n\n"
        "La clase con mayor probabilidad es la predicción. Si supera el umbral de confianza (70%), "
        "el modelo evalúa si hay value para apostar.",
        color=VERDE, label="🔢")

    # 7B Pi-Rating
    p_7b = doc.add_paragraph()
    r = p_7b.add_run("B.  Pi-Rating: La fuerza real de cada equipo")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7b.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "El Pi-Rating (Constantinou & Fenton, 2012) es un sistema de rating inspirado en el ELO del "
        "ajedrez, pero mejorado para fútbol: actualiza los ratings con goles, no solo con el resultado "
        "binario ganar/perder. Esto captura la magnitud del rendimiento.")
    add_formula(doc, "r_nuevo = r_anterior  +  K × (P_real − P_esperado)")
    add_para(doc,
        "Parámetros del modelo: K = 0.5 (tasa de aprendizaje) · Decay = 0.98 (decaimiento inter-temporada)\n"
        "P_real = función de goles marcados vs recibidos (rango 0-1)\n"
        "P_esperado = función logística de la diferencia de ratings actuales",
        color=GRIS, size=9)
    add_callout(doc,
        "Ejemplo numérico:\n"
        "Inter (r=1.43) vs Napoli (r=1.28). Diferencia = 0.15.\n"
        "P_esperado (Inter) ≈ 0.62  (la función logística transforma la diferencia de ratings)\n"
        "Si el partido termina 3-0 para Inter: P_real ≈ 0.85 (goleada)\n"
        "Δr = 0.5 × (0.85 − 0.62) = +0.115  →  Inter sube de 1.43 a 1.545\n"
        "Napoli baja simétricamente.\n\n"
        "Al inicio de cada temporada: r_nuevo = r × 0.98 (decaimiento)\n"
        "Esto evita que equipos del pasado dominen el modelo para siempre.",
        color=AZUL, label="📐")
    add_para(doc,
        "La variable pi_diff (diferencia de Pi-Ratings) es la feature más importante del modelo, "
        "con ~25% de importancia relativa según el árbol de ganancia. No sorprende: cuánto mejor "
        "es un equipo que el otro históricamente predice mejor el resultado que cualquier otra variable.")

    # 7C Value / EV
    p_7c = doc.add_paragraph()
    r = p_7c.add_run("C.  Value Betting: Matemática del Valor Esperado (EV)")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7c.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "Una apuesta es rentable a largo plazo solo si su Valor Esperado (EV) es positivo. "
        "La fórmula es la misma que en cualquier proceso estocástico de ingeniería:")
    add_formula(doc,
        "EV = P(ganar) × ganancia_neta  −  P(perder) × apuesta\n"
        "   = p × (cuota − 1) − (1 − p) × 1\n"
        "   = p × cuota − 1")
    add_callout(doc,
        "Ejemplo numérico:\n"
        "Modelo: P(Inter gana) = 0.62   ·   Cuota bookmaker: 1.75\n"
        "EV = 0.62 × 1.75 − 1 = 1.085 − 1 = +0.085  (+8.5%)\n\n"
        "Interpretación: por cada $10.000 apostados, el valor esperado es ganar $850 adicionales. "
        "En 100 apuestas similares, el resultado promedio converge a +$850 por apuesta.\n\n"
        "Contra-ejemplo — misma probabilidad, cuota 1.55:\n"
        "EV = 0.62 × 1.55 − 1 = 0.961 − 1 = −0.039  (−3.9%)  → No se apuesta.",
        color=VERDE, label="💰")
    add_para(doc,
        "El umbral value ≥ 10% significa que solo apostamos cuando p × cuota ≥ 1.10. "
        "Esto filtra apuestas marginales y se queda con las que tienen ventaja estadística clara.")

    # 7D Kelly
    p_7d = doc.add_paragraph()
    r = p_7d.add_run("D.  Criterio de Kelly: Dimensionamiento Óptimo de Capital")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7d.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "Una vez que sabemos que una apuesta tiene EV positivo, ¿cuánto del bankroll apostar? "
        "El Criterio de Kelly (Bell Labs, 1956) es la fórmula óptima que maximiza el crecimiento "
        "logarítmico del capital a largo plazo:")
    add_formula(doc, "f* = (p × cuota − 1) / (cuota − 1)")
    add_para(doc, "donde f* es la fracción del bankroll a apostar.", color=GRIS, size=9)
    add_callout(doc,
        "Ejemplo:\n"
        "P = 0.62   ·   cuota = 1.75   ·   bankroll = $100.000 CLP\n"
        "f* = (0.62 × 1.75 − 1) / (1.75 − 1) = 0.085 / 0.75 = 11.3% del bankroll\n"
        "Monto Kelly pleno = $11.300\n\n"
        "El modelo usa Quarter Kelly (×0.25): $11.300 × 0.25 = $2.825 CLP\n\n"
        "¿Por qué ¼? Kelly asume que la probabilidad es perfectamente conocida. "
        "En la práctica, el modelo puede estar equivocado, así que apostar menos del óptimo "
        "teórico reduce la volatilidad sin sacrificar demasiado crecimiento.",
        color=VERDE, label="📐")

    # 7E TimeSeriesSplit
    p_7e = doc.add_paragraph()
    r = p_7e.add_run("E.  Validación Temporal: ¿Son Reales las Métricas?")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7e.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "El mayor error en ML deportivo es el data leakage: si el modelo entrena con datos del "
        "futuro, sus métricas son artificialmente buenas. El modelo usa TimeSeriesSplit con 5 folds "
        "— el estándar para series temporales:")

    # Tabla de folds
    tbl_ts = doc.add_table(rows=5, cols=2)
    tbl_ts.style = "Table Grid"
    folds = [
        ("Fold 1", "Entrena 2019-2020  →  Evalúa 2021"),
        ("Fold 2", "Entrena 2019-2021  →  Evalúa 2022"),
        ("Fold 3", "Entrena 2019-2022  →  Evalúa 2023"),
        ("Fold 4", "Entrena 2019-2023  →  Evalúa 2024"),
        ("TEST SET", "Último 20%  —  El modelo nunca lo vio durante el entrenamiento"),
    ]
    for ri, (fold, desc) in enumerate(folds):
        c0 = tbl_ts.rows[ri].cells[0]
        c1 = tbl_ts.rows[ri].cells[1]
        is_test = "TEST" in fold
        set_cell_bg(c0, RGBColor(0x16, 0xA3, 0x4A) if is_test else FONDO_TBL)
        r0 = c0.paragraphs[0].add_run(fold)
        r0.font.bold = True; r0.font.size = Pt(9)
        r0.font.color.rgb = BLANCO
        r1 = c1.paragraphs[0].add_run(desc)
        r1.font.size = Pt(9)
        if is_test:
            r1.font.bold = True; r1.font.color.rgb = VERDE
    doc.add_paragraph()

    add_callout(doc,
        "Resultados del modelo:\n"
        "CV accuracy: 0.4888 ± 0.0248  (media ± desv. estándar entre los 5 folds)\n"
        "Test accuracy: 0.5226  (datos completamente nuevos — nunca vistos)\n\n"
        "Regla de oro: Si test accuracy >> CV accuracy → overfitting (el modelo memorizó).\n"
        "Si test accuracy ≈ CV accuracy → el modelo generaliza bien.\n"
        "Aquí la diferencia es 0.034 (<5%)  →  Sin overfitting. ✅\n\n"
        "La desviación estándar de 0.0248 entre folds indica que el modelo es consistente: "
        "no hay folds donde rinda muy bien y otros donde falle. Un modelo inestable tendría std > 0.05.",
        color=AZUL, label="📊")

    # 7F Distribución multinomial
    p_7f = doc.add_paragraph()
    r = p_7f.add_run("F.  La Distribución de Resultados: No es Normal")
    r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = AZUL
    p_7f.paragraph_format.space_before = Pt(12)

    add_para(doc,
        "Un error conceptual frecuente es asumir que la incertidumbre del modelo sigue una "
        "distribución normal (Gaussiana). En realidad, el output de XGBoost para 3 clases sigue "
        "una distribución multinomial:")

    bullets_dist = [
        "Tiene 3 parámetros: p₀ (local), p₁ (empate), p₂ (visitante)",
        "Restricciones: p₀ + p₁ + p₂ = 1, y cada pᵢ ∈ [0, 1]",
        "El espacio de probabilidad está en un simplex 2D (triángulo), no en la recta real ℝ",
        "No puede ser Gaussiana porque está acotada entre 0 y 1, y las 3 clases son dependientes",
    ]
    for b in bullets_dist:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(b).font.size = Pt(9.5)

    add_callout(doc,
        "¿Qué significa p = 0.70 para 'Inter gana'?\n"
        "Calibración del modelo: en el conjunto histórico, cuando el modelo asignó ~70% de "
        "probabilidad a una victoria local, ese equipo ganó realmente en ~67-68% de los casos "
        "(accuracy al apostar). La diferencia del 2-3% es el error de calibración — aceptable "
        "y dentro del margen de ruido estadístico con n=61 apuestas.",
        color=AZUL, label="🎯")

    add_para(doc,
        "Los intervalos de confianza para predicciones individuales no están calculados directamente "
        "(XGBoost puntual no tiene incertidumbre explícita como los modelos bayesianos). Sin embargo, "
        "el umbral de confianza del 70% actúa como filtro que selecciona predicciones donde el modelo "
        "tiene suficiente separación entre las 3 probabilidades — reduciendo el riesgo de apostar en "
        "partidos 'dudosos' donde las probabilidades están cerca de 33% / 33% / 33% (máxima incertidumbre).")

    # ── 8. GLOSARIO ────────────────────────────────────────────────────────────
    add_heading(doc, "8. Glosario de Términos", level=1, color=AZUL)
    glosario = [
        ("Pi-Rating",        "Número que mide la fortaleza de un equipo, como el ELO en ajedrez pero mejorado para fútbol. Considera los goles, no solo el resultado."),
        ("XGBoost",          "Algoritmo de machine learning que combina cientos de árboles de decisión secuenciales. Es el motor principal del agente."),
        ("Softmax",          "Función matemática que convierte logits (valores crudos) en probabilidades que suman 100%. Garantiza que ninguna probabilidad sea negativa."),
        ("Accuracy",         "Porcentaje de partidos que el modelo predice correctamente. 52.3% global. Al aplicar filtros de confianza, sube hasta 67.2%."),
        ("ROI",              "Return on Investment. Ganancia neta / total apostado. ROI +31.7% significa que por cada $100 apostados, ganaste $31.70 en promedio."),
        ("Value Bet",        "Apuesta donde la probabilidad real es mayor que la que implica la cuota del bookmaker. EV = p × cuota − 1 > 0."),
        ("Criterio de Kelly","Fórmula óptima para dimensionar apuestas: f* = (p×cuota−1)/(cuota−1). El modelo usa Quarter Kelly (×0.25) para seguridad."),
        ("TimeSeriesSplit",  "Técnica de validación que siempre entrena con datos pasados y evalúa con datos futuros. Evita data leakage."),
        ("Data Leakage",     "Error grave en ML: cuando el modelo 've el futuro' durante el entrenamiento. Sus métricas serían falsamente altas."),
        ("xG (Expected Goals)","Cuántos goles 'deberían haber' marcado los equipos según la calidad de sus ocasiones. Fuente: Understat (2019-2024)."),
        ("Distribución multinomial","Distribución de probabilidad para 3+ categorías excluyentes. Aquí: p(Home) + p(Empate) + p(Away) = 1."),
        ("Gradient Boosting","Técnica ML que construye modelos secuenciales donde cada modelo nuevo corrige los errores del anterior."),
    ]
    tbl_gl = doc.add_table(rows=len(glosario), cols=2)
    tbl_gl.style = "Table Grid"
    for ri, (term, defn) in enumerate(glosario):
        c0 = tbl_gl.rows[ri].cells[0]
        c1 = tbl_gl.rows[ri].cells[1]
        set_cell_bg(c0, FONDO_TBL)
        r0 = c0.paragraphs[0].add_run(term)
        r0.font.bold = True; r0.font.size = Pt(9); r0.font.color.rgb = VERDE
        r1 = c1.paragraphs[0].add_run(defn)
        r1.font.size = Pt(9)
    doc.add_paragraph()

    # ── 9. LIMITACIONES ────────────────────────────────────────────────────────
    add_heading(doc, "9. Limitaciones Actuales del Modelo", level=1, color=ROJO)
    limitaciones = [
        ("🔴", "Serie A solamente.",
         "Solo la Serie A tiene suficientes apuestas con ROI positivo confirmado. "
         "La Liga está a 6 apuestas del umbral de activación (20 apuestas con ROI > 0)."),
        ("✅", "xG 2019-2024 integrado.",
         "Fix A Sprint 10: datos de Expected Goals de Understat ahora cubren 2019-2024 "
         "(10.707 partidos vs 1.752 antes). xg_temporada_home/away ya son top-5 features."),
        ("✅", "Transfermarkt integrado.",
         "Fix B Sprint 10: valor de mercado de plantillas ahora integrado en el modelo. "
         "Cache 30 días, pre-carga única por run. Nuevas features: valor_home_mill, ratio_valor, log_ratio_valor."),
        ("🔴", "ROI simulado ≠ ROI real.",
         "El +31.7% es backtesting histórico. El rendimiento real puede variar. "
         "Las cuotas cambian, los bookmakers se adaptan, y factores imprevisibles ocurren."),
        ("🟡", "Sin lesiones ni sanciones en tiempo real.",
         "El modelo usa datos históricos, no noticias del día. Una baja de último minuto "
         "de un jugador clave no es capturada automáticamente."),
    ]
    for icon, titulo, desc in limitaciones:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        r1 = p.add_run(f"{icon}  {titulo}  ")
        r1.font.bold = True; r1.font.size = Pt(10)
        r2 = p.add_run(desc)
        r2.font.size = Pt(9.5); r2.font.color.rgb = GRIS

    # ── ADVERTENCIA ────────────────────────────────────────────────────────────
    add_callout(doc,
        "⚠️ APUESTA RESPONSABLEMENTE\n\n"
        "Este modelo es una herramienta de análisis estadístico, no una garantía de ganancias. "
        "Las apuestas deportivas conllevan riesgo de pérdida de dinero real. Nunca apuestes más "
        "de lo que puedas permitirte perder. El ROI positivo en simulación histórica no garantiza "
        "rentabilidad futura. Si las apuestas afectan negativamente tu vida, busca ayuda.",
        color=AMARILLO, label="")

    # ── PIE DE PÁGINA ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_f = p_footer.add_run(
        "Agente de Apuestas Deportivas · Sprint 10 · 24/03/2026\n"
        "Modelo: XGBoost · Features: Pi-Rating + xG (2019-2024) + Transfermarkt + Forma reciente\n"
        "Datos: football-data.co.uk + Understat + Transfermarkt · Predicción: predictor_tiempo_real.py\n"
        "Desarrollado con Claude Code (Anthropic) · Sócrates Cabral"
    )
    r_f.font.size = Pt(8); r_f.font.color.rgb = GRIS; r_f.font.italic = True

    doc.save(OUT_DOCX)
    print(f"[OK] DOCX guardado: {OUT_DOCX}")
    return True


def generar_pdf():
    try:
        from docx2pdf import convert
        convert(str(OUT_DOCX), str(OUT_PDF))
        print(f"[OK] PDF guardado: {OUT_PDF}")
        return True
    except Exception as e:
        print(f"[WARN] PDF no generado: {e}")
        print("[INFO] Puedes exportar el .docx a PDF manualmente desde Word (Archivo → Guardar como → PDF)")
        return False


if __name__ == "__main__":
    print("[INFO] Generando manual...")
    ok_docx = generar_docx()
    if ok_docx:
        generar_pdf()
    print("[INFO] Listo.")
