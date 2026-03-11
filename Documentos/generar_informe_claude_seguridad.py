"""
Genera: Informe_Claude_Egakat_Seguridad_TI.docx
Carpeta: C:\ClaudeWork\Documentos\
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime
import os

OUTPUT_PATH = r"C:\ClaudeWork\Documentos\Informe_Claude_Egakat_Seguridad_TI.docx"

# ─── Colores corporativos ───────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1F, 0x35, 0x64)   # Encabezados principales
AZUL_MEDIO   = RGBColor(0x2E, 0x6E, 0xB4)   # Subtitulos
NARANJA      = RGBColor(0xE0, 0x6C, 0x00)   # Alertas / destacados
VERDE        = RGBColor(0x1A, 0x7A, 0x3C)   # Positivo
GRIS_TEXTO   = RGBColor(0x3C, 0x3C, 0x3C)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

# ─── Helpers ────────────────────────────────────────────────────
def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(16)
    run.font.color.rgb = AZUL_OSCURO

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(13)
    run.font.color.rgb = AZUL_MEDIO

def heading3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11)
    run.font.color.rgb = NARANJA

def body(doc, text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTO
    return p

def bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = GRIS_TEXTO

def tabla_header(table, headers, bg="1F3564"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = h
        set_cell_bg(cell, bg)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.color.rgb = BLANCO
            run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

def tabla_fila(table, valores, bg=None, color=None):
    row = table.add_row()
    for i, v in enumerate(valores):
        cell = row.cells[i]
        cell.text = v
        if bg:
            set_cell_bg(cell, bg)
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(10)
            if color:
                run.font.color.rgb = color

def separador(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run("─" * 80)
    run.font.size  = Pt(8)
    run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

# ─── DOCUMENTO ──────────────────────────────────────────────────
doc = Document()

# Margenes
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── PORTADA ─────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("\n\n\nEGAKAT SPA — Control Management & Continuous Improvement")
run.bold = True; run.font.size = Pt(13); run.font.color.rgb = AZUL_MEDIO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("INTELIGENCIA ARTIFICIAL EN OPERACIONES\nAsistente Claude — Evaluacion de Seguridad, Productividad y Gobernanza")
run.bold = True; run.font.size = Pt(20); run.font.color.rgb = AZUL_OSCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"\nDocumento Tecnico para Equipo de TI\nFecha: {datetime.date.today().strftime('%d de %B de %Y')}\nAutor: Socrates Cabral — Head of Control Management & Continuous Improvement")
run.font.size = Pt(11); run.font.color.rgb = GRIS_TEXTO

doc.add_page_break()

# ── 1. RESUMEN EJECUTIVO ─────────────────────────────────────────
heading1(doc, "1. Resumen Ejecutivo")
body(doc, (
    "Este documento presenta el analisis completo del uso del asistente de inteligencia artificial "
    "Claude (Anthropic) en la gestion operacional de Egakat SPA. Describe las herramientas implementadas, "
    "los beneficios de productividad obtenidos, el marco de seguridad aplicado, y los argumentos tecnicos "
    "y de negocio para su uso controlado en equipos corporativos."
))
body(doc, (
    "Claude no es un acceso irrestricto a internet ni una herramienta de riesgo descontrolado. "
    "Es un asistente especializado que opera localmente, bajo reglas definidas, con acceso "
    "delimitado unicamente a las carpetas y sistemas autorizados por el administrador."
))

separador(doc)

# ── 2. QUE ES CLAUDE ─────────────────────────────────────────────
heading1(doc, "2. Que es Claude y Como Funciona")
heading2(doc, "2.1 Descripcion General")
body(doc, (
    "Claude es un modelo de lenguaje de gran escala (LLM) desarrollado por Anthropic, empresa de "
    "seguridad en IA fundada en 2021 por ex-investigadores de OpenAI. A diferencia de ChatGPT (OpenAI) "
    "o Gemini (Google), Anthropic enfoca su desarrollo en IA segura, interpretable y alineada con valores humanos."
))

heading2(doc, "2.2 Modalidades de Uso")
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Modalidad", "Descripcion", "Acceso a archivos locales"])
tabla_fila(t, ["Claude.ai (web/desktop)", "Chat conversacional, analisis de documentos, diseno de logica", "No — solo lo que se pega en el chat"])
tabla_fila(t, ["Claude Code (terminal)", "Agente de codigo: lee, edita y ejecuta en la maquina local", "Si — controlado por permisos del usuario"])
tabla_fila(t, ["MCP Servers", "Extensiones que dan capacidades adicionales (navegador, archivos, BI)", "Segun configuracion — directorios especificos"])

separador(doc)

# ── 3. HERRAMIENTAS INSTALADAS ────────────────────────────────────
heading1(doc, "3. Herramientas Instaladas en Egakat")
heading2(doc, "3.1 Claude Code — Agente de Automatizacion")
body(doc, (
    "Claude Code es la interfaz de linea de comandos (CLI) de Claude. Opera directamente en la terminal "
    "del equipo y puede leer archivos, ejecutar scripts Python, buscar codigo y editar archivos — "
    "siempre bajo supervision del usuario que aprueba cada accion."
))
for b in [
    "Desarrolla y corrige scripts Python de automatizacion WMS",
    "Lee logs de errores y propone soluciones en segundos",
    "Gestiona la memoria del proyecto entre sesiones (MEMORY.md)",
    "Ejecuta comandos del sistema (schtasks, powershell) con aprobacion explicita",
]:
    bullet(doc, b)

heading2(doc, "3.2 MCP Server — Playwright (Control de Navegador)")
body(doc, (
    "Playwright MCP permite que Claude controle un navegador web directamente desde la sesion de chat. "
    "Esto habilita el debug visual del sistema WMS sin necesidad de describir errores con palabras."
))
for b in [
    "Navega al WMS y toma screenshots en tiempo real",
    "Inspecciona elementos de la pagina para diagnosticar fallos de automatizacion",
    "Ejecuta acciones en el navegador (clicks, formularios, descargas) bajo supervision",
    "El navegador se abre en la maquina local — no hay datos enviados a servidores externos",
]:
    bullet(doc, b)

heading2(doc, "3.3 MCP Server — Filesystem (Acceso a Archivos)")
body(doc, (
    "Filesystem MCP otorga a Claude acceso de lectura/escritura a directorios especificamente autorizados. "
    "Fuera de esos directorios, Claude no puede ver ni tocar ningun archivo."
))
for b in [
    "Directorios autorizados: C:\\ClaudeWork\\ y OneDrive Egakat",
    "Claude puede leer reportes Excel, logs y scripts sin que el usuario los copie manualmente",
    "Acceso denegado automaticamente a: C:\\Users\\, documentos personales, credenciales del sistema",
    "Ninguna escritura ocurre sin que Claude informe al usuario la accion que realizara",
]:
    bullet(doc, b)

heading2(doc, "3.4 MCP Server — Power BI (Modelo de Datos)")
body(doc, (
    "Conecta Claude con el modelo semantico de Power BI Desktop instalado en el equipo. "
    "Permite consultar medidas, tablas y relaciones del modelo sin abrir la interfaz grafica."
))

separador(doc)

# ── 4. PRODUCTIVIDAD Y HORAS AHORRADAS ───────────────────────────
heading1(doc, "4. Productividad y Horas Hombre Ahorradas")
heading2(doc, "4.1 Procesos Automatizados con Claude")

t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
tabla_header(t, ["Proceso", "Tiempo Manual (estimado)", "Tiempo Automatizado", "Ahorro Semanal"])
tabla_fila(t, ["Descarga Stock WMS (3 centros)", "45 min/dia", "0 min (Task Scheduler 8AM)", "3.75 hrs"])
tabla_fila(t, ["Consulta de Posiciones (8 reportes)", "30 min/dia", "0 min (automatico)", "2.5 hrs"])
tabla_fila(t, ["Staging IN/OUT (16 clientes)", "60 min/dia", "0 min (automatico)", "5 hrs"])
tabla_fila(t, ["Comparador VDR Derco Parts", "90 min/revision", "0 min (cada hora automatico)", "6+ hrs"])
tabla_fila(t, ["Debug de errores en scripts", "2-4 hrs/incidente", "15-30 min con Claude Code", "Variable"])
tabla_fila(t, ["TOTAL ESTIMADO", "", "", "~17-20 hrs/semana"], bg="1F3564", color=BLANCO)

heading2(doc, "4.2 Velocidad de Desarrollo")
body(doc, "El desarrollo de scripts con Claude como asistente reduce drasticamente los ciclos de prueba y error:")
for b in [
    "Script nuevo desde cero: de 2-3 dias a 4-8 horas (reduccion del 70%)",
    "Correccion de bug critico: de horas a minutos (Claude lee el log y propone el fix directo)",
    "Documentacion tecnica: de 4-6 horas a 30 minutos (Claude genera el borrador, el usuario revisa)",
    "Configuracion de Task Scheduler + Power Automate: de investigacion propia (dias) a guia paso a paso",
]:
    bullet(doc, b)

heading2(doc, "4.3 Impacto en el Equipo")
body(doc, (
    "El tiempo liberado (~17-20 hrs/semana) se redirige a analisis de datos, mejora continua de procesos "
    "y proyectos de mayor valor agregado — en lugar de tareas operativas repetitivas y descarga manual de reportes."
))

separador(doc)

# ── 5. SEGURIDAD ─────────────────────────────────────────────────
heading1(doc, "5. Analisis de Seguridad")
heading2(doc, "5.1 Arquitectura de Datos — Que Sale del Equipo y Que No")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Herramienta", "Datos que salen al exterior", "Datos que permanecen locales"])
tabla_fila(t, ["Claude.ai (chat)", "Lo que el usuario pega manualmente en el chat", "Todo lo demas — archivos, credenciales, base de datos"])
tabla_fila(t, ["Claude Code (terminal)", "Solo el texto del chat (sin archivos adjuntos automaticos)", "Archivos, scripts, logs, .env — todo local"])
tabla_fila(t, ["MCP Playwright", "NADA — el navegador corre en la maquina local", "Todo: sesiones WMS, descargas, formularios"])
tabla_fila(t, ["MCP Filesystem", "NADA — lee archivos localmente y los muestra en el chat", "Archivos Excel, CSVs, scripts — no se suben"])
tabla_fila(t, ["MCP Power BI", "NADA — consulta el modelo local de PBI Desktop", "Modelo semantico, medidas, relaciones"])

heading2(doc, "5.2 Politica de Anthropic sobre Datos")
for b in [
    "Anthropic NO entrena sus modelos con conversaciones de usuarios de planes de pago (Team/Enterprise)",
    "Los datos de conversacion se pueden eliminar desde la cuenta de usuario",
    "Anthropic cumple con SOC 2 Type II, ISO 27001 en proceso de certificacion",
    "Politica de privacidad disponible en: anthropic.com/privacy",
    "Los datos en transito se cifran con TLS 1.2/1.3",
]:
    bullet(doc, b)

heading2(doc, "5.3 Riesgos Identificados y Mitigaciones")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Riesgo", "Nivel", "Mitigacion Aplicada"])
tabla_fila(t, ["Usuario pega credenciales en el chat", "ALTO", "Regla: jamas pegar passwords ni tokens en Claude.ai — usar .env local"])
tabla_fila(t, ["Acceso a archivos fuera del scope", "MEDIO", "Filesystem MCP limitado a C:\\ClaudeWork\\ y OneDrive Egakat unicamente"])
tabla_fila(t, ["Claude ejecuta comandos destructivos", "MEDIO", "Claude Code requiere aprobacion explicita del usuario para cada accion"])
tabla_fila(t, ["Filtracion de datos de clientes", "ALTO", "No pegar datos de clientes (RUT, contratos, precios) en Claude.ai"])
tabla_fila(t, ["Acceso no autorizado a la sesion", "BAJO", "El equipo tiene BitLocker + credenciales Windows corporativas"])
tabla_fila(t, ["Uso por terceros no autorizados", "BAJO", "Claude Code vinculado a cuenta personal de Anthropic — 1 usuario"])

separador(doc)

# ── 6. PROS Y CONTRAS ─────────────────────────────────────────────
heading1(doc, "6. Pros y Contras de la Implementacion")

heading2(doc, "6.1 Ventajas")
for b in [
    "Automatizacion de tareas repetitivas sin necesidad de contratar programador externo",
    "Reduccion de errores humanos en procesos de descarga y comparacion de datos",
    "Disponibilidad 24/7 — los scripts corren automaticamente aunque el usuario no este presente",
    "Curva de aprendizaje baja — el usuario guia a Claude en lenguaje natural",
    "Memoria persistente del proyecto — no se pierde contexto entre sesiones",
    "Acceso controlado y auditable — cada accion queda registrada en logs locales",
    "Sin instalacion de servidores propios — Anthropic gestiona la infraestructura del modelo",
    "Integracion con herramientas existentes: OneDrive, Power BI, Task Scheduler, Power Automate",
]:
    bullet(doc, "+" + " " + b)

heading2(doc, "6.2 Desventajas y Limitaciones")
for b in [
    "Costo mensual de la suscripcion Claude (plan Team: USD 25/usuario/mes aprox.)",
    "Requiere conexion a internet para el componente de chat (Claude.ai / Claude Code)",
    "El modelo puede cometer errores — toda salida debe ser revisada por el usuario",
    "No reemplaza al personal: requiere supervision humana para validar resultados criticos",
    "Riesgo de uso inapropiado si el usuario no sigue las politicas de datos definidas",
    "Contexto de sesion limitado — conversaciones muy largas pueden perder detalles anteriores",
]:
    bullet(doc, "-" + " " + b)

separador(doc)

# ── 7. ARGUMENTOS ANTE TI ─────────────────────────────────────────
heading1(doc, "7. Argumentos ante el Equipo de TI")

heading2(doc, "7.1 'Por que instalar una IA en un equipo corporativo?'")
body(doc, "Respuesta: Claude no se 'instala' en el sentido tradicional. Es una aplicacion de escritorio estandar (similar a Microsoft Teams o Zoom) que se comunica via HTTPS con servidores de Anthropic. No modifica el sistema operativo, no instala drivers, no abre puertos de red.")

heading2(doc, "7.2 'Que pasa con los datos de la empresa?'")
body(doc, (
    "Solo se transmiten al servidor de Anthropic los textos que el usuario escribe en el chat. "
    "Los archivos locales (Excel, CSV, scripts) nunca se suben automaticamente. "
    "El MCP Filesystem lee archivos localmente y presenta el contenido en la sesion — "
    "funcionalmente equivalente a que el usuario copie y pegue el texto."
))

heading2(doc, "7.3 'Es Claude mas riesgoso que Google o Microsoft Copilot?'")
body(doc, (
    "No. Microsoft 365 Copilot tiene acceso completo a todos los correos, archivos de SharePoint, "
    "Teams y OneDrive del tenant. Claude tiene acceso unicamente a lo que el usuario le muestra "
    "explicitamente o a las carpetas configuradas en MCP Filesystem. "
    "En terminos de superficie de exposicion, Claude es significativamente mas restringido."
))

heading2(doc, "7.4 'Que control tiene TI sobre esto?'")
for b in [
    "Pueden revisar el archivo de configuracion MCP en cualquier momento (ruta documentada en este informe)",
    "Pueden auditar los logs de ejecucion en C:\\ClaudeWork\\logs\\",
    "Pueden bloquear dominios de Anthropic en el firewall si es necesario",
    "Pueden solicitar que se eliminen directorios del MCP Filesystem para reducir el scope de acceso",
    "La cuenta de Anthropic esta asociada al correo corporativo del usuario — rastreable",
]:
    bullet(doc, b)

heading2(doc, "7.5 Comparativa con Herramientas Corporativas Aprobadas")
t = doc.add_table(rows=1, cols=4)
t.style = "Table Grid"
tabla_header(t, ["Herramienta", "Acceso a datos corporativos", "Datos salen al exterior", "Control de TI"])
tabla_fila(t, ["Microsoft 365 Copilot", "TODO el tenant (correo, SharePoint, Teams)", "Si — servidores Microsoft", "Politica de tenant"])
tabla_fila(t, ["Google Workspace AI", "TODO el workspace (Drive, Gmail, Meet)", "Si — servidores Google", "Politica de dominio"])
tabla_fila(t, ["Claude (configuracion actual)", "Solo C:\\ClaudeWork\\ y OneDrive Egakat", "Solo texto del chat", "Archivo MCP config + logs locales"])
tabla_fila(t, ["ChatGPT Plus (sin config)", "Solo lo que el usuario pega", "Si — servidores OpenAI", "Ninguno"])

separador(doc)

# ── 8. RESTRICCIONES DE SEGURIDAD RECOMENDADAS ────────────────────
heading1(doc, "8. Restricciones de Seguridad a Implementar")

heading2(doc, "8.1 Restricciones Inmediatas (Hoy)")
for b in [
    "REGLA 1 — Nunca pegar en Claude: passwords, tokens API, claves .env, datos de tarjetas, RUT de personas naturales",
    "REGLA 2 — Nunca pegar en Claude: contratos con clientes, precios negociados, margenes, informacion financiera confidencial",
    "REGLA 3 — Nunca pegar en Claude: datos personales de empleados (sueldos, evaluaciones, informacion medica)",
    "REGLA 4 — El archivo .env nunca debe incluirse en el scope del MCP Filesystem",
    "REGLA 5 — Revisar siempre el codigo generado antes de ejecutarlo en produccion",
    "REGLA 6 — No aprobar acciones destructivas de Claude Code sin entender que hace (rm, drop, delete)",
]:
    bullet(doc, b)

heading2(doc, "8.2 Configuraciones Tecnicas Recomendadas")
for b in [
    "Mantener BitLocker activo en el equipo — si se pierde el equipo, los datos estan cifrados",
    "Usar cuenta Windows con contrasena fuerte — Claude Code hereda los permisos del usuario activo",
    "No ejecutar Claude Code como Administrador — usar cuenta de usuario estandar",
    "Revisar periodicamente los directorios en MCP Filesystem y eliminar los que ya no apliquen",
    "Activar bloqueo automatico de pantalla (maximo 5 minutos de inactividad)",
    "No compartir la sesion de Claude con otros usuarios — cada persona debe tener su propia cuenta",
]:
    bullet(doc, b)

heading2(doc, "8.3 Opciones de Seguridad Adicional (Avanzadas)")
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Medida", "Descripcion", "Dificultad"])
tabla_fila(t, ["Firewall de aplicacion", "Bloquear Claude a nivel de red si el equipo no esta en uso", "Baja — configuracion de Windows Defender Firewall"])
tabla_fila(t, ["Audit log de MCP", "Registrar cada llamada al MCP Filesystem en un log separado", "Media — requiere wrapper en el servidor MCP"])
tabla_fila(t, ["Directorio MCP de solo lectura", "Configurar OneDrive como read-only en el MCP", "Baja — modificar args en claude_desktop_config.json"])
tabla_fila(t, ["DLP (Data Loss Prevention)", "Politica de Microsoft Purview que bloquee dominios Anthropic si detecta datos sensibles", "Alta — requiere licencia Microsoft 365 E3+"])
tabla_fila(t, ["VPN corporativa obligatoria", "Exigir VPN activa para usar Claude Code — el trafico pasa por el proxy corporativo", "Media — politica de red"])

heading2(doc, "8.4 Clasificacion de Datos — Guia Rapida")

t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Tipo de dato", "Se puede usar con Claude?", "Como usarlo"])
tabla_fila(t, ["Nombres de clientes (empresa)", "SI", "Sin restriccion — dato publico"])
tabla_fila(t, ["Stock WMS / posiciones", "SI", "Sin restriccion — dato operacional interno"])
tabla_fila(t, ["RUT de empresas clientes", "SI con precaucion", "Solo para scripts — no compartir masivamente"])
tabla_fila(t, ["Precios / tarifas contractuales", "NO en Claude.ai", "Solo en Claude Code local si es necesario"])
tabla_fila(t, ["Datos de empleados", "NO", "Nunca — Ley 19.628 Chile (datos personales)"])
tabla_fila(t, ["Credenciales / passwords", "NUNCA", "Siempre en .env — nunca en el chat"])
tabla_fila(t, ["Informacion financiera Egakat", "NO en Claude.ai", "Consultar con Gerencia antes de usar"])

separador(doc)

# ── 9. PLAN DE GOBERNANZA ─────────────────────────────────────────
heading1(doc, "9. Plan de Gobernanza Propuesto")

body(doc, "Para operar con Claude de forma responsable y auditable en Egakat SPA, se propone el siguiente marco minimo de gobernanza:")

heading2(doc, "9.1 Roles y Responsabilidades")
t = doc.add_table(rows=1, cols=3)
t.style = "Table Grid"
tabla_header(t, ["Rol", "Responsabilidad", "Frecuencia"])
tabla_fila(t, ["Usuario (Socrates Cabral)", "Operar Claude, aprobar acciones, no compartir datos sensibles", "Diaria"])
tabla_fila(t, ["TI (jcontreras@tinetservices.cl)", "Revisar configuracion MCP, auditar logs, aprobar cambios de scope", "Mensual"])
tabla_fila(t, ["Gerencia / Control Management", "Validar que los reportes generados son correctos y confiables", "Semanal"])

heading2(doc, "9.2 Revision Periodica")
for b in [
    "Mensual: revisar directorios configurados en MCP Filesystem — eliminar los que ya no apliquen",
    "Trimestral: auditar logs de ejecucion en C:\\ClaudeWork\\logs\\ buscando anomalias",
    "Semestral: revisar politica de datos con TI y actualizar segun nuevos proyectos",
    "Anual: evaluar alternativas de IA y comparar con la solucion actual",
]:
    bullet(doc, b)

separador(doc)

# ── 10. CONCLUSION ────────────────────────────────────────────────
heading1(doc, "10. Conclusion")
body(doc, (
    "La implementacion de Claude en el equipo de Control Management & Continuous Improvement de Egakat SPA "
    "representa una ventaja operacional significativa: automatizacion de procesos repetitivos, reduccion de "
    "errores, y liberacion de tiempo para tareas de mayor valor estrategico."
))
body(doc, (
    "Desde el punto de vista de seguridad, la configuracion actual es conservadora y controlada: "
    "los datos no salen del equipo salvo lo que el usuario escribe explicitamente en el chat, "
    "y el acceso a archivos esta limitado a directorios especificos autorizados. "
    "Esta configuracion es, en terminos de exposicion de datos, mas segura que Microsoft 365 Copilot "
    "o Google Workspace AI, que tienen acceso completo al tenant corporativo."
))
body(doc, (
    "Con las restricciones de seguridad documentadas en este informe y el plan de gobernanza propuesto, "
    "el uso de Claude puede formalizarse como una herramienta de productividad corporativa auditada "
    "y bajo control de TI."
))

# Firma
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(
    f"Socrates Cabral\nHead of Control Management & Continuous Improvement\nEgakat SPA\n{datetime.date.today().strftime('%d/%m/%Y')}"
)
run.font.size = Pt(10)
run.font.color.rgb = GRIS_TEXTO
run.bold = True

# ── GUARDAR ───────────────────────────────────────────────────────
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
doc.save(OUTPUT_PATH)
print(f"Documento generado: {OUTPUT_PATH}")
