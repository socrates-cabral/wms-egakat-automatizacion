"""
Genera: Analisis_Encuesta_NPS_Egakat_2026.docx
Ubicacion: C:\ClaudeWork\Documentos\
"""
import sys
sys.stdout.reconfigure(encoding="utf-8")

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Colores corporativos ──────────────────────────────────────────────────────
AZUL_OSCURO  = RGBColor(0x1F, 0x49, 0x7D)   # azul corporativo
AZUL_MEDIO   = RGBColor(0x2E, 0x75, 0xB6)
GRIS_OSCURO  = RGBColor(0x40, 0x40, 0x40)
GRIS_CLARO   = RGBColor(0xF2, 0xF2, 0xF2)
VERDE        = RGBColor(0x37, 0x86, 0x44)
ROJO         = RGBColor(0xC0, 0x00, 0x00)
NARANJA      = RGBColor(0xED, 0x7D, 0x31)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

doc = Document()

# ── Margenes ──────────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_cell_bg(cell, color_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  color_hex)
    tcPr.append(shd)

def set_cell_borders(cell, color="BFBFBF", sz="4"):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top","left","bottom","right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    sz)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)

def heading1(text, color=AZUL_OSCURO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(14)
    run.font.color.rgb = color
    # linea inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "1F497D")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text, color=AZUL_MEDIO):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold      = True
    run.font.size = Pt(11.5)
    run.font.color.rgb = color
    return p

def body(text, bold_parts=None, color=None, size=10.5, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0, color=GRIS_OSCURO):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.5 + level * 0.5)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = color
    return p

def callout(text, bg="EAF4FB", border_color="2E75B6"):
    """Caja de texto destacada (tabla 1x1)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cell = tbl.cell(0, 0)
    cell.width = Inches(6)
    set_cell_bg(cell, bg)
    set_cell_borders(cell, color=border_color, sz="8")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.left_indent  = Cm(0.3)
    run = p.add_run(text)
    run.font.size = Pt(10.5)
    run.font.color.rgb = GRIS_OSCURO
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def add_table(headers, rows_data, header_bg="1F497D", alt_bg="EBF3FB"):
    cols = len(headers)
    tbl  = doc.add_table(rows=1+len(rows_data), cols=cols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Encabezado
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        set_cell_bg(cell, header_bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = BLANCO

    # Filas
    for r_idx, row_data in enumerate(rows_data):
        row = tbl.rows[r_idx + 1]
        bg  = alt_bg if r_idx % 2 == 0 else "FFFFFF"
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(10)
            run.font.color.rgb = GRIS_OSCURO

    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

# =============================================================================
# PORTADA
# =============================================================================
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("EGAKAT SPA")
run.bold = True
run.font.size = Pt(13)
run.font.color.rgb = AZUL_MEDIO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Analisis y Propuesta de Mejora")
run.bold = True
run.font.size = Pt(22)
run.font.color.rgb = AZUL_OSCURO

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Encuesta NPS 2026")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = AZUL_MEDIO

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Instrumento de Medicion de Satisfaccion y Lealtad de Clientes")
run.font.size = Pt(12)
run.font.color.rgb = GRIS_OSCURO
run.italic = True

doc.add_paragraph()
doc.add_paragraph()

# Caja portada
tbl = doc.add_table(rows=4, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
info = [
    ("Preparado por:", "Control Management & Continuous Improvement"),
    ("Fecha:",         datetime.date.today().strftime("%d de %B de %Y")),
    ("Clasificacion:", "Uso interno — Gerencia"),
    ("Version:",       "1.0"),
]
for i, (label, value) in enumerate(info):
    c0 = tbl.rows[i].cells[0]
    c1 = tbl.rows[i].cells[1]
    set_cell_bg(c0, "1F497D")
    set_cell_bg(c1, "EBF3FB" if i % 2 == 0 else "FFFFFF")
    set_cell_borders(c0, "FFFFFF")
    set_cell_borders(c1, "BDD7EE")
    r0 = c0.paragraphs[0].add_run(label)
    r0.bold = True; r0.font.size = Pt(10); r0.font.color.rgb = BLANCO
    r1 = c1.paragraphs[0].add_run(value)
    r1.font.size = Pt(10); r1.font.color.rgb = GRIS_OSCURO

doc.add_page_break()

# =============================================================================
# 1. RESUMEN EJECUTIVO
# =============================================================================
heading1("1. Resumen Ejecutivo")
callout(
    "La encuesta actual combina tres metodologias distintas (NPS, CSAT y CES) en un solo "
    "instrumento de 16 preguntas. Si bien esto genera datos ricos para operaciones, no es "
    "optimo para reportar un indicador NPS unico y accionable a nivel gerencial. "
    "Este documento propone separar el instrumento en dos encuestas complementarias y "
    "establece el framework de calculo, interpretacion y seguimiento del NPS.",
    bg="EAF4FB", border_color="2E75B6"
)

body("Los puntos clave de este analisis son:", size=10.5)
bullet("La pregunta NPS (C1) usa escala 1-10 en vez del estandar internacional 0-10, afectando el calculo.")
bullet("La encuesta de 16 preguntas reduce la tasa de respuesta y diluye el indicador principal.")
bullet("Se propone un modelo de dos instrumentos: NPS trimestral (3 preguntas) + CSAT semestral (16 preguntas).")
bullet("Se definen KPIs, benchmarks de industria 3PL y un plan de implementacion.")

# =============================================================================
# 2. QUE ES EL NPS Y POR QUE IMPORTA
# =============================================================================
heading1("2. Que es el NPS y por que importa para Egakat")

heading2("2.1 Definicion")
body(
    "El Net Promoter Score (NPS) es un indicador de lealtad del cliente creado por "
    "Fred Reichheld (Bain & Company, 2003) y adoptado globalmente como el estandar "
    "para medir la probabilidad de recomendacion. Se basa en una sola pregunta:"
)
callout(
    '"En una escala de 0 a 10, que tan probable es que recomiendes a Egakat '
    'a un colega de negocios o amigo?"',
    bg="FFF2CC", border_color="ED7D31"
)

heading2("2.2 Clasificacion de respuestas")
add_table(
    ["Puntaje", "Categoria", "Descripcion", "Accion recomendada"],
    [
        ["9 - 10", "Promotores",  "Clientes leales que recomiendan activamente",     "Fidelizar y usar como referencia"],
        ["7 - 8",  "Pasivos",     "Satisfechos pero vulnerables a la competencia",   "Mejorar para convertir en promotores"],
        ["0 - 6",  "Detractores", "Insatisfechos que pueden daniar la reputacion",   "Contactar, escuchar y resolver urgente"],
    ]
)

heading2("2.3 Formula de calculo")
callout(
    "NPS = % Promotores - % Detractores\n\n"
    "Ejemplo: 60 respuestas -> 30 promotores (50%), 10 detractores (17%), 20 pasivos (33%)\n"
    "NPS = 50% - 17% = +33",
    bg="E2EFDA", border_color="378644"
)

heading2("2.4 Interpretacion del resultado")
add_table(
    ["Rango NPS", "Evaluacion", "Significado"],
    [
        ["70 a 100",  "Excelente",  "Liderazgo en lealtad de clientes"],
        ["50 a 69",   "Muy bueno",  "Alta satisfaccion, mejora continua posible"],
        ["30 a 49",   "Bueno",      "Mayoria satisfecha, oportunidades claras"],
        ["0 a 29",    "Regular",    "Mas trabajo necesario, riesgo de fuga"],
        ["-100 a -1", "Critico",    "Mas detractores que promotores — accion inmediata"],
    ]
)

# =============================================================================
# 3. DIAGNOSTICO DE LA ENCUESTA ACTUAL
# =============================================================================
heading1("3. Diagnostico de la Encuesta Actual")

heading2("3.1 Estructura actual")
add_table(
    ["Seccion", "Preguntas", "Metodologia", "Proposito"],
    [
        ["A - Experiencia con Egakat",  "A1 a A7",  "CSAT",       "Satisfaccion operacional por area"],
        ["B - Resolucion de problemas", "B1 a B4",  "CES",        "Esfuerzo en resolucion de incidencias"],
        ["C - Tu opinion nos importa",  "C1 a C5",  "NPS + datos","Recomendacion y datos de contacto"],
    ]
)

heading2("3.2 Problemas identificados")

body("Problema 1: Escala NPS incorrecta", bold_parts=["Problema 1"])
p = doc.add_paragraph()
run = p.add_run("Problema 1 — Escala NPS incorrecta")
run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = ROJO
bullet("La pregunta C1 usa escala 1-10 en lugar del estandar 0-10.")
bullet("Impacto: los detractores deberian incluir puntajes 0-6; al iniciar en 1 se pierde un punto de la escala y los benchmarks no son comparables.")
bullet("Solucion: cambiar a 0-10 en LimeSurvey.")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Problema 2 — La pregunta NPS esta enterrada en la pregunta 13 de 16")
run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = ROJO
bullet("Si el encuestado abandona la encuesta antes de llegar a C1, no se obtiene el NPS.")
bullet("El NPS es el KPI principal. Debe ir primero o en posicion destacada.")
bullet("Solucion: en la encuesta NPS corta, moverla a la primera pregunta.")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Problema 3 — Ausencia de logica condicional en Seccion B")
run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NARANJA
bullet("Si B1 = 'No', las preguntas B2-B4 son irrelevantes pero igual se muestran.")
bullet("Esto genera friction y puede provocar abandono.")
bullet("Solucion: configurar skip logic en LimeSurvey (si B1 != Si, saltar a Seccion C).")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Problema 4 — Mezcla de metodologias sin separacion clara")
run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NARANJA
bullet("NPS, CSAT y CES miden cosas distintas y tienen diferentes frecuencias optimas.")
bullet("Mezclarlos en una encuesta anual reduce la utilidad de cada indicador.")
bullet("Solucion: separar en dos instrumentos (ver Seccion 5).")

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Problema 5 — Datos de contacto insuficientes para segmentacion")
run.bold = True; run.font.size = Pt(10.5); run.font.color.rgb = NARANJA
bullet("C5 solicita Nombre, Telefono y Correo, pero no Empresa ni Cargo.")
bullet("Sin Empresa no se puede segmentar el NPS por cliente ni por industria.")
bullet("Solucion: agregar campos Empresa y Cargo en C5.")

heading2("3.3 Fortalezas de la encuesta actual")
bullet("Estructura logica y flujo bien organizado en tres secciones.", color=VERDE)
bullet("Inclusion de preguntas abiertas (A3, A4, C3) para capturar insights cualitativos.", color=VERDE)
bullet("Seccion B sobre resolucion de problemas es valiosa y diferenciadora — pocas empresas 3PL la miden.", color=VERDE)
bullet("Escalas descriptivas (Muy mal/Mal/Aceptable/Bien/Muy bien) facilitan la respuesta.", color=VERDE)
bullet("Preguntas A5-A7 sobre SLA, precision y visibilidad son relevantes y accionables.", color=VERDE)
bullet("Redaccion clara, tono apropiado y tiempo estimado razonable.", color=VERDE)

# =============================================================================
# 4. BENCHMARKS INDUSTRIA 3PL
# =============================================================================
heading1("4. Benchmarks de NPS en Logistica y 3PL")

body(
    "Conocer los benchmarks del sector permite contextualizar el resultado de Egakat "
    "y establecer metas realistas para la gerencia."
)
add_table(
    ["Sector / Empresa",          "NPS Promedio",  "Fuente / Referencia"],
    [
        ["Logistica y 3PL (global)",    "40 - 50",  "Bain & Company, 2024"],
        ["Supply Chain / Freight",      "35 - 45",  "Satmetrix Industry Benchmark"],
        ["Operadores logisticos LAT",   "30 - 45",  "CustomerGauge LatAm Report"],
        ["DHL Supply Chain",            "~62",       "DHL Group Annual Report"],
        ["FedEx Supply Chain",          "~55",       "Satmetrix 2023"],
        ["Meta interna sugerida Y1",    "> 40",      "Propuesta Egakat 2026"],
    ]
)
callout(
    "Recomendacion: establecer la meta NPS 2026 en >= 40, con revision trimestral. "
    "El primer resultado servira como baseline y permitira definir metas mas precisas para 2027.",
    bg="E2EFDA", border_color="378644"
)

# =============================================================================
# 5. PROPUESTA: DOS INSTRUMENTOS
# =============================================================================
heading1("5. Propuesta: Modelo de Dos Instrumentos")

body(
    "Se propone reemplazar la encuesta unica por dos instrumentos complementarios, "
    "cada uno con un proposito, frecuencia y audiencia definidos."
)

add_table(
    ["",                  "Encuesta NPS",                              "Encuesta CSAT Operacional"],
    [
        ["Proposito",     "KPI de lealtad para gerencia",              "Diagnostico operacional por area"],
        ["Preguntas",     "3",                                          "16"],
        ["Tiempo",        "< 1 minuto",                                 "3-5 minutos"],
        ["Frecuencia",    "Trimestral o post-servicio",                 "Semestral"],
        ["Audiencia",     "Todos los clientes activos",                 "Clientes clave (muestra)"],
        ["Canal",         "Email automatico / WhatsApp",                "Email con aviso previo"],
        ["KPI principal", "NPS score",                                  "CSAT por area + CES problemas"],
        ["Reporte",       "Dashboard gerencial mensual",                "Informe semestral operaciones"],
    ]
)

# ── Encuesta NPS corta ────────────────────────────────────────────────────────
heading2("5.1 Propuesta: Encuesta NPS (version corta — 3 preguntas)")

callout(
    "Esta encuesta es el instrumento principal para el indicador gerencial. "
    "Debe ser breve, simple y enviarse con alta frecuencia para obtener datos estadisticamente validos.",
    bg="EAF4FB", border_color="2E75B6"
)

add_table(
    ["N°", "Pregunta", "Tipo", "Escala"],
    [
        ["NPS1", "En una escala de 0 a 10, que tan probable es que recomiendes los servicios de Egakat a un colega o amigo?", "NPS",    "0-10 (0=Nada probable, 10=Extremadamente probable)"],
        ["NPS2", "En una o dos palabras, cual es la principal razon de tu puntuacion?",                                        "Abierta","Texto libre (max 150 caracteres)"],
        ["NPS3", "Te gustaria que nos contactemos contigo para conocer mejor tu experiencia?",                                 "Si/No",  "Si -> captura nombre, empresa, cargo, email"],
    ]
)

body("Notas de configuracion en LimeSurvey:", size=10)
bullet("NPS1: tipo 'Array (Numbers)' o 'Multiple choice' con opciones 0 al 10 en una sola fila.")
bullet("NPS2: tipo 'Long free text', opcional.")
bullet("NPS3: tipo 'Yes/No' con condicion: si = Si -> mostrar grupo de preguntas de contacto.")
bullet("Activar anonimato: Si para NPS2 y NPS3 opcionales para maximizar respuestas honestas.")

# ── Encuesta CSAT ─────────────────────────────────────────────────────────────
heading2("5.2 Encuesta CSAT Operacional (version actual mejorada)")

body(
    "La encuesta actual (16 preguntas) es adecuada para el diagnostico operacional "
    "con los siguientes ajustes:"
)
add_table(
    ["Pregunta", "Cambio recomendado",                                              "Prioridad"],
    [
        ["A1",  "Mantener redaccion actual (ya corregida)",                           "OK"],
        ["A2",  "Mantener 'Aceptable' (ya corregida)",                                "OK"],
        ["A5-A7","Mantener preguntas de SLA, precision y visibilidad",                "OK"],
        ["B1",  "Agregar skip logic: si No/No recuerdo -> saltar a Seccion C",        "Alta"],
        ["C1",  "Mover a Encuesta NPS corta. En CSAT omitir o dejar como secundaria", "Alta"],
        ["C5",  "Agregar campos Empresa y Cargo",                                     "Alta"],
        ["C2",  "Configurar seleccion multiple con maximo 3 opciones",               "Media"],
    ]
)

# =============================================================================
# 6. PLAN DE IMPLEMENTACION
# =============================================================================
heading1("6. Plan de Implementacion")

add_table(
    ["Fase", "Accion",                                                         "Responsable",          "Plazo"],
    [
        ["1", "Corregir escala C1 a 0-10 en encuesta actual",                   "Control Management",   "Semana 1"],
        ["2", "Crear encuesta NPS corta (3 preguntas) en LimeSurvey",           "Control Management",   "Semana 1-2"],
        ["3", "Agregar skip logic en Seccion B",                                 "Control Management",   "Semana 2"],
        ["4", "Agregar campos Empresa y Cargo en C5",                            "Control Management",   "Semana 2"],
        ["5", "Primer envio NPS a todos los clientes activos",                   "Comercial + CC",       "Semana 3"],
        ["6", "Establecer automatizacion de envio trimestral (LimeSurvey/email)","TI / Control",         "Mes 1"],
        ["7", "Construir dashboard NPS en Power BI",                             "Control Management",   "Mes 2"],
        ["8", "Primer reporte NPS a gerencia con baseline 2026",                 "Control Management",   "Mes 2"],
        ["9", "Envio CSAT Operacional semestral (instrumento completo)",         "Control Management",   "Mes 6"],
    ]
)

# =============================================================================
# 7. KPIs PARA DASHBOARD GERENCIAL
# =============================================================================
heading1("7. KPIs Sugeridos para Dashboard Gerencial")

body(
    "El reporte gerencial debe ser simple, visual y accionable. "
    "Se recomienda un dashboard mensual con los siguientes indicadores:"
)

add_table(
    ["KPI",                    "Descripcion",                                         "Frecuencia", "Meta 2026"],
    [
        ["NPS Score",          "Net Promoter Score global",                            "Mensual",    ">= 40"],
        ["% Promotores",       "Clientes con puntaje 9-10",                            "Mensual",    ">= 50%"],
        ["% Detractores",      "Clientes con puntaje 0-6",                             "Mensual",    "<= 15%"],
        ["Tasa de respuesta",  "% clientes que completaron la encuesta NPS",           "Mensual",    ">= 30%"],
        ["NPS por cliente",    "Score individual por empresa cliente",                 "Trimestral", "N/A"],
        ["CSAT promedio",      "Promedio satisfaccion areas A2 (1-5)",                 "Semestral",  ">= 4.0"],
        ["CES problemas",      "% clientes: resolvemos 'siempre' o 'casi siempre'",   "Semestral",  ">= 70%"],
        ["Problemas reportados","% clientes que indicaron problema en B1",             "Semestral",  "<= 20%"],
    ]
)

heading2("7.1 Segmentacion recomendada del NPS")
body("Para mayor profundidad, segmentar el NPS por:")
bullet("Cliente / empresa")
bullet("Centro de distribucion (Quilicura, Pudahuel, Pudahuel Unitario)")
bullet("Linea de servicio (Recepcion, Despacho, Inventario, Transporte)")
bullet("Periodo (comparativa trimestral y anual)")

# =============================================================================
# 8. ACCIONES SEGUN CATEGORIA NPS
# =============================================================================
heading1("8. Protocolo de Accion por Categoria NPS")

body(
    "Para que el NPS genere valor real, debe existir un protocolo de seguimiento "
    "diferenciado segun la categoria del cliente:"
)

add_table(
    ["Categoria",    "Puntuacion", "Tiempo de respuesta",  "Accion"],
    [
        ["Promotor",    "9-10",   "72 horas",   "Enviar agradecimiento. Solicitar testimonio o caso de exito. Considerar para referidos."],
        ["Pasivo",      "7-8",    "5 dias",     "Identificar que mejoraria su experiencia. Seguimiento por ejecutivo de cuenta."],
        ["Detractor",   "0-6",    "24 horas",   "Contacto directo por jefatura o gerencia. Escuchar, documentar y crear plan de mejora con fecha."],
    ]
)

callout(
    "Regla de oro: ninguna respuesta de detractor debe quedar sin atencion en mas de 24 horas. "
    "Un detractor bien atendido puede convertirse en promotor — y un detractor ignorado "
    "puede convertirse en un cliente perdido y una resena negativa publica.",
    bg="FFF2CC", border_color="ED7D31"
)

# =============================================================================
# 9. RECOMENDACIONES FINALES
# =============================================================================
heading1("9. Recomendaciones Finales")

add_table(
    ["#",  "Recomendacion",                                                                "Impacto",  "Esfuerzo"],
    [
        ["1", "Implementar encuesta NPS corta (3 preguntas) con escala 0-10",              "Alto",     "Bajo"],
        ["2", "Enviar NPS trimestral a todos los clientes activos",                        "Alto",     "Bajo"],
        ["3", "Configurar skip logic en Seccion B de encuesta CSAT",                       "Medio",    "Bajo"],
        ["4", "Agregar Empresa y Cargo en datos de contacto",                              "Medio",    "Bajo"],
        ["5", "Construir dashboard NPS en Power BI con los KPIs definidos",                "Alto",     "Medio"],
        ["6", "Definir protocolo de atencion a detractores (< 24 horas)",                  "Alto",     "Medio"],
        ["7", "Establecer baseline NPS 2026 con primer envio masivo",                      "Alto",     "Bajo"],
        ["8", "Segmentar NPS por cliente y centro de distribucion",                        "Alto",     "Medio"],
        ["9", "Revisar y actualizar preguntas CSAT semestralmente segun feedback interno", "Medio",    "Bajo"],
    ]
)

# =============================================================================
# CIERRE
# =============================================================================
doc.add_paragraph()
callout(
    "Conclusion: La encuesta actual es un buen instrumento de diagnostico operacional. "
    "Con ajustes menores y la separacion en dos instrumentos, Egakat podra reportar "
    "un NPS confiable, comparable con benchmarks de industria y accionable para la "
    "gerencia — convirtiendo la voz del cliente en una ventaja competitiva real.",
    bg="E2EFDA", border_color="378644"
)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run(f"Egakat SPA | Control Management & Continuous Improvement | {datetime.date.today().strftime('%B %Y')}")
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
run.italic = True

# ── Guardar ───────────────────────────────────────────────────────────────────
output = r"C:\ClaudeWork\Documentos\Analisis_Encuesta_NPS_Egakat_2026.docx"
doc.save(output)
print(f"Documento generado: {output}")
