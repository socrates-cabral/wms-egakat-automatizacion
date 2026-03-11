"""
Genera el documento de respuesta a José Contreras (TI) sobre App Registration Azure AD.
"""

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

doc = Document()

# Estilos base
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Encabezado
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
run = p.add_run("Santiago, 10 de marzo de 2026")
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Para: José Contreras — TI Tinet Services\n")
run.bold = True
run = p.add_run("De: Sócrates Cabral — Control Management & Mejora Continua, Egakat\n")
run = p.add_run("Asunto: App Registration Azure AD — información adicional solicitada")

doc.add_paragraph()

# Cuerpo
doc.add_paragraph(
    "Hola José, gracias por revisar el tema y por el feedback. Aquí te paso lo que pediste."
)

doc.add_paragraph()

# Sección 1
p = doc.add_paragraph()
run = p.add_run("Documentación oficial de Microsoft")
run.bold = True
run.font.size = Pt(11)

doc.add_paragraph(
    "Para el registro de la app y la autenticación sin usuario (client credentials flow), "
    "la referencia principal está acá:"
)

refs = [
    ("Registrar una aplicación en Azure AD (Entra ID):",
     "https://learn.microsoft.com/es-es/entra/identity-platform/quickstart-register-app"),
    ("Configurar acceso de aplicación a SharePoint (App-Only):",
     "https://learn.microsoft.com/es-es/sharepoint/dev/solution-guidance/security-apponly-azureacs"),
    ("Flujo OAuth2 client credentials (sin usuario interactivo):",
     "https://learn.microsoft.com/es-es/azure/active-directory/develop/v2-oauth2-client-creds-grant-flow"),
    ("Permiso Sites.Selected (acceso acotado a un sitio específico):",
     "https://learn.microsoft.com/es-es/graph/permissions-reference#sitesselected"),
]

for titulo, url in refs:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(titulo + "\n")
    run.bold = True
    run = p.add_run(url)
    run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
    run.font.size = Pt(10)

doc.add_paragraph()

# Sección 2 — permisos
p = doc.add_paragraph()
run = p.add_run("Sobre los permisos — no se necesita admin global")
run.bold = True

doc.add_paragraph(
    "Entiendo que el punto de preocupación es darle perfil de admin de SharePoint a la app. "
    "En realidad eso no es necesario. El script usa el permiso Sites.Selected, que es la "
    "opción que Microsoft recomienda justamente para este tipo de integraciones: le das acceso "
    "únicamente al sitio que necesitas, no a todo el tenant."
)

doc.add_paragraph(
    "Los dos permisos de aplicación que se necesitan son:"
)

permisos = [
    ("Sites.Selected", "Acceso solo al sitio específico de SharePoint, nada más."),
    ("Files.ReadWrite", "Para subir y leer archivos en la librería de documentos del sitio."),
]

for perm, desc in permisos:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(perm + ": ")
    run.bold = True
    p.add_run(desc)

doc.add_paragraph()

# Sección 3 — script
p = doc.add_paragraph()
run = p.add_run("El script")
run.bold = True

doc.add_paragraph(
    "Te adjunto el archivo sharepoint_upload.py. Para resumir lo que hace: se conecta a Azure AD "
    "usando client_id y client_secret (los genera Azure al registrar la app), y sube los reportes "
    "Excel al sitio https://egakatcom.sharepoint.com/sites/DatosparaDashboard. "
    "Las credenciales no están en el código, las lee desde un archivo .env local en mi equipo."
)

doc.add_paragraph(
    "Lo que necesitaría de TI es solo esto:"
)

pasos = [
    "Registrar la app en Azure AD Entra y obtener el client_id y client_secret.",
    "Conceder el consentimiento de administrador para los permisos Sites.Selected y Files.ReadWrite en el sitio específico.",
    "Pasarme las credenciales generadas a socrates.cabral@egakat.cl.",
]

for paso in pasos:
    p = doc.add_paragraph(style='List Number')
    p.add_run(paso)

doc.add_paragraph()

doc.add_paragraph(
    "Si necesitas que revisemos algo juntos o tienes dudas del script, con gusto coordinamos "
    "una llamada corta. Quedo atento."
)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Sócrates Cabral\n")
run.bold = True
p.add_run("Control Management & Mejora Continua\n")
p.add_run("Egakat SPA")

# Guardar
salida = Path(r"C:\ClaudeWork\Solicitudes_IT\Respuesta_Jose_Contreras_AppRegistration.docx")
doc.save(salida)
print(f"Documento generado: {salida}")
